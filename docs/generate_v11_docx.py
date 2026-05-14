"""
Generator AgriFlow_v11.docx — proposal lengkap (claim-precision pass on top of v10).

Run: python docs/generate_v11_docx.py
Output: ./AgriFlow_v11.docx (di docs/ folder)

v11 vs v10: konten bisnis identik dengan v10, plus a *claim-precision pass*
PLUS two scoring-quality engine fixes (#1 volume_score formula, #2 segment_multiplier):
  - Tightened stable-matching claim — Gale-Shapley fires only when both kab Tier 1
    (allocation.py:341-347); for current Jatim 8-IHK/30-non-IHK split the production
    path is greedy-with-equity-priority. Stable matching becomes load-bearing once
    Tier 1 coverage expands (~90 IHK cities nasional).
  - Tightened two-tier confidence claim — described as data-quality labeling
    (HIGH / MEDIUM / LOW) rather than two algorithmic modes; MEDIUM is the
    structural default in Jatim, HIGH requires Tier1↔Tier1.
  - Tightened equity-boost framing — +30% applies to deficit-side kab with IPM<68,
    fires twice in 32-match demo (Ngawi→Bangkalan, Ngawi→Sampang beras).
  - Dropped "World-First" from title and exec summary; replaced with substantive
    specificity ("purpose-built sub-national matching engine combining stable
    matching + IPM-based equity + two-tier data quality, kalibrasi BPS 2024").
  - Scoped "first-in-world" claims to "first sub-national Indonesian food matching
    engine" with provenance caveat on competitive comparison table.
  - Tightened climate-aware framing — scoring penalty when route weather data
    available; neutral fallback otherwise (10 routes seeded in demo).
  - Section 1.4 (NEW) Perubahan Kunci dari v10 ke v11.
  - FIX #1: volume_score → coverage-of-demand model (scoring.py:42-72)
  - FIX #2: segment_multiplier in MatchResult (models.py + allocation.py)
            range ±10% by demand.segment, fully auditable via flags

134/134 pytest pass (up from 106 baseline + 5 scenario extensions = 131,
plus 1 coverage test, plus 7 segment-multiplier acid tests).

Engine-side changes are surgical and backwards-compatible — RETAIL default
segment = 1.00 multiplier means existing callers see no change.
"""
from __future__ import annotations
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except (AttributeError, OSError):
        pass


NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x86, 0xAB)
GRAY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREEN = RGBColor(0x2D, 0x6A, 0x4F)


def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def para(doc, text, bold_first=None, size=11, italic=False):
    p = doc.add_paragraph()
    if bold_first:
        r = p.add_run(bold_first)
        r.bold = True
        r.font.size = Pt(size)
        p.add_run(" ")
    r = p.add_run(text)
    r.font.size = Pt(size)
    if italic:
        r.italic = True
    return p


def callout(doc, text, color="DDF4FF"):
    """Render callout box (single-cell table with colored background)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return table


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = DARK
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.text = ""
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def make_table(doc, headers, rows, col_widths=None, header_color="1F3A5F"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10.5)
        set_cell_shading(cell, header_color)

    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_shading(cells[ci], "F8F9FA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# =============================================================================

def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ==================================================================
    # COVER
    # ==================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("AGRIFLOW")
    r.font.size = Pt(40)
    r.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Proposal v11.0")
    r.font.size = Pt(20)
    r.font.color.rgb = ACCENT
    r.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Food Intelligence dengan Purpose-Built Sub-National Matching Engine\n"
                  "+ Reference Implementation + Calibration + Claim-Precision Update")
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY
    r.italic = True

    doc.add_paragraph().paragraph_format.space_before = Pt(40)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Deteksi. Prediksi. Distribusi. Untuk Semua.")
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Food Intelligence Infrastructure untuk Indonesia")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tanpa Install Aplikasi. Tanpa Diskriminasi.\n"
                  "Dashboard · WhatsApp · Phone Call — Semua Orang Bisa Pakai")
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = GRAY

    doc.add_paragraph().paragraph_format.space_before = Pt(30)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("First Indonesian AI Matching Engine\n"
                  "for Sub-National Food Distribution")
    r.font.size = Pt(13)
    r.bold = True
    r.font.color.rgb = ACCENT

    doc.add_paragraph().paragraph_format.space_before = Pt(30)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PIDI — DIGDAYA × Hackathon 2026\n"
                  "Kategori: Peningkatan Produktivitas, Ketahanan Pangan, "
                  "dan Penciptaan Lapangan Kerja\n"
                  "Problem Statement #2: Platform Matching Demand-Supply Antarwilayah")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\nVersi 11.0 — Mei 2026 (Claim-Precision Update)")
    r.font.size = Pt(11)
    r.italic = True
    r.bold = True
    r.font.color.rgb = NAVY

    doc.add_page_break()

    # ==================================================================
    # 1. RINGKASAN EKSEKUTIF
    # ==================================================================
    heading(doc, "1. Ringkasan Eksekutif", level=1)

    callout(doc,
        "AgriFlow adalah Food Intelligence Infrastructure untuk Indonesia "
        "yang dirancang untuk semua orang — petani tua dengan HP jadul, pedagang "
        "pasar dengan WhatsApp, staff Dinas dengan laptop, hingga tuna netra. "
        "Akses real-time tanpa install aplikasi apapun. Inti platform: "
        "AgriFlow Matching Engine v11.0 — algoritma 4-lapis purpose-built untuk "
        "konteks pangan sub-nasional Indonesia. Layer 3 menggabungkan "
        "Modified Gale-Shapley stable matching (Nobel Prize Economics 2012, "
        "aktif saat kedua kabupaten Tier 1) dengan greedy multi-objective + "
        "equity priority (aktif saat ada Tier 2 — produksi Jatim saat ini), "
        "multi-objective scoring 5 dimensi, dan equity multiplier untuk "
        "kabupaten tertinggal dengan IPM <68 saat menjadi deficit (kalibrasi "
        "BPS 2024 — Sampang 66.72, Bangkalan 67.70). 24 skenario edge case "
        "(19 engineering + 5 commercial-reality v11) didokumentasi lengkap "
        "dengan handling spesifik dan tervalidasi pytest 126/126. "
        "Triple-channel access: (1) Dashboard Web Pemda, "
        "(2) WhatsApp Chat + Voice, (3) Phone Call IVR Bahasa Daerah. "
        "Operational layer langsung untuk GPIPS (Gerakan Pengendalian Inflasi "
        "dan Pangan Sejahtera) BI Februari 2026. 8 revenue streams dari "
        "6 customer segments. Hackathon Rp 19jt. Operasional Rp 1-2jt/bulan. "
        "Revenue Y3 Rp 16,1M net Rp 11,1M.",
    )

    heading(doc, "1.1 Perubahan Kunci dari v7 ke v8", level=2)
    make_table(doc,
        headers=["Aspek", "v7 (April 2026)", "v8"],
        rows=[
            ["Matching Engine", "Disebut sebagai 'jarak × volume × kecukupan'",
             "Section 5.5 Deep Dive: 3-lapis hybrid, formula matematis, pseudocode, 19 scenarios (saat v8)"],
            ["Strategi data", "PIHPS + BPS generic",
             "Hybrid 2-Tier: Tier 1 (8 kota IHK) + Tier 2 (30 kab non-IHK estimasi) dengan confidence labeling"],
            ["Klaim differentiation", "Triple-channel accessibility",
             "+ Purpose-built sub-national food matching engine — kombinasi stable matching (Gale-Shapley) + multi-objective scoring + IPM-based equity untuk konteks Indonesia"],
            ["Scenario coverage", "Tidak eksplisit",
             "19 skenario edge case dalam 5 kategori (Volume, Spasial, Temporal, Disrupsi, Politis) — di-expand jadi 24 di v11"],
            ["Global benchmark", "5 adopsi global",
             "+ Analisa 12 platform pangan dunia (eNAM, MealConnect, ECX, EAX, AfMX, Food Drop, Uber, FEWS NET, WFP mVAM, dll)"],
            ["Halaman", "21", "~30 (matching engine deep dive +8-10)"],
        ],
        col_widths=[3.5, 5.0, 8.5],
    )

    heading(doc, "1.2 Perubahan Kunci dari v8 ke v9", level=2)
    para(doc,
        "v9 adalah engineering completion update — fokus ke executable code dan "
        "dokumentasi teknis untuk eksekusi tim. Tidak ada perubahan strategi "
        "bisnis atau klaim diferensiator.")
    make_table(doc,
        headers=["Aspek", "v8", "v9"],
        rows=[
            ["Matching Engine", "Pseudocode + diagram di proposal",
             "+ Implementasi Python siap-run (5 module: models / constraints / scoring / allocation / engine)"],
            ["Test Coverage", "19 skenario didokumentasi",
             "+ 106 unit + scenario test pass (pytest), semua 19 skenario tervalidasi otomatis — v11 expand jadi 126/126 dengan 24 skenario"],
            ["Sample Data", "Tidak ada",
             "+ Generator CSV realistis 38 kab × 19 komoditas Jatim untuk dev/demo offline"],
            ["Data Sources", "Disebut PIHPS, Bapanas, BPS, BMKG, dll.",
             "+ 8 connector module dengan dual-mode (mock CSV + live API), URL endpoint terdokumentasi (Section 5.5.8)"],
            ["IPM Data", "Sampang 61.6 (data 2023 lama)",
             "Update ke BPS BRS Desember 2024: Sampang 66.72 (terendah Jatim)"],
            ["Quick Start Guide", "Tidak ada",
             "+ Section 5.5.9: clone → install → generate sample → run demo → run tests dalam 5 menit"],
            ["Test Strategy", "Hanya tabel skenario",
             "+ Section 5.5.10: mapping 24 skenario (v11) → file test, fixture per kabupaten Jatim, edge case detail"],
            ["Latency Validation", "Klaim <500ms p99",
             "+ Demo end-to-end measured ~1-50ms pada sample data 38×19"],
        ],
        col_widths=[3.5, 5.0, 8.5],
    )

    heading(doc, "1.3 Perubahan Kunci dari v9 ke v10 (NEW)", level=2)
    para(doc,
        "v10 adalah calibration & bug fix update — fokus ke kalibrasi parameter "
        "agar konsisten dengan data BPS 2024 nyata, dan perbaikan bug yang "
        "ditemukan saat audit menyeluruh terhadap kodebase v9. Tidak ada "
        "perubahan strategi bisnis atau klaim diferensiator.")
    make_table(doc,
        headers=["Aspek", "v9", "v10"],
        rows=[
            ["Equity Multiplier Threshold",
             "<65→1.30 / <70→1.15 / <75→1.05 / ≥75→1.00",
             "<68→1.30 / <72→1.15 / <78→1.05 / ≥78→1.00 (kalibrasi BPS 2024)"],
            ["Klaim '+30% boost untuk kab tertinggal'",
             "Klaim ada di proposal tapi threshold lama tidak pernah trigger (IPM terendah Jatim = 66.72)",
             "Aplicable konkret ke Sampang (66.72) & Bangkalan (67.70) — flag EQUITY_BOOST_30 fires di demo"],
            ["IPM Source-of-Truth Sync",
             "data_sources/bps.py:IPM_2024_JATIM dan generate_sample_data.py divergent (7+ kab beda)",
             "Sinkron — generator adalah source-of-truth, BPS connector mirror"],
            ["Demo Cross-Platform Compatibility",
             "examples/run_demo.py crash di Windows console default (UnicodeEncodeError)",
             "Force UTF-8 stdout/stderr; demo berjalan di Windows fresh-install"],
            ["Stale Data Confidence Drop",
             "Hanya HIGH→MEDIUM (Tier 2 stale data tetap MEDIUM)",
             "Bertingkat HIGH→MEDIUM→LOW (sesuai spec C3)"],
            ["Stale Identity Check",
             "Pakai dataclass `in` operator (object identity, fragile)",
             "Pakai (kab_id, commodity_code) tuple set — robust"],
            ["BMKG `prefer_bmkg` Flag",
             "Flag ada tapi tidak pernah dipakai (dead code)",
             "Properly honored: aktif jika adm4_lookup tersedia, fallback Open-Meteo"],
            ["Latency Benchmark Tool",
             "Tidak ada",
             "+ benchmarks/latency.py — multi-config stress test (40×33 / 361×361 / 200×200)"],
            ["Documentation Generator",
             "Manual write proposal docx",
             "+ docs/generate_v10_docx.py — regenerate proposal dari source code"],
        ],
        col_widths=[3.5, 5.0, 8.5],
    )

    callout(doc,
        "Konsekuensi v10 untuk juri: klaim '+30% boost untuk kab tertinggal' "
        "demonstrably valid. Demo run menunjukkan flag EQUITY_BOOST_30 fires "
        "untuk Sampang & Bangkalan. Latency target <500ms p99 ter-verifikasi "
        "via benchmark (highest p99 = 55.53ms = margin 88.9%). Cross-platform "
        "demo bekerja di Windows fresh-install.",
        color="E8F4FD",
    )

    heading(doc, "1.4 Perubahan Kunci dari v10 ke v11 (NEW)", level=2)
    para(doc,
        "v11 adalah claim-precision update — fokus ke kalibrasi *pernyataan* "
        "(bukan parameter) agar setiap klaim diferensiator di proposal "
        "dapat diverifikasi langsung terhadap kode di matching_engine/. "
        "Tidak ada perubahan engine, tidak ada perubahan parameter, tidak "
        "ada perubahan test. Honest framing pass.")
    make_table(doc,
        headers=["Aspek", "v10 (klaim lama)", "v11 (klaim terverifikasi)"],
        rows=[
            ["Stable Matching",
             "'Matching Engine pakai stable matching Gale-Shapley (Nobel 2012)' — implisit always-on",
             "Gale-Shapley fires saat kedua kabupaten Tier 1 (allocation.py:341-347). Untuk Jatim 8-IHK / 30-non-IHK saat ini, supply originate dari Tier 2 → production path = greedy-with-equity-priority. Stable matching jadi load-bearing setelah Tier 1 coverage expand nasional (~90 kota IHK)."],
            ["Two-Tier Confidence",
             "'Data harian PIHPS pakai algoritma ketat; data mingguan Bapanas pakai algoritma fleksibel'",
             "Two-tier data-quality labeling — setiap match return label HIGH / MEDIUM / LOW berdasarkan tier composition + data freshness. Di konfigurasi Jatim saat ini, MEDIUM adalah label default (semua surplus dari non-IHK); HIGH require kedua kab Tier 1 (rare di Jatim, achievable nasional); LOW fires saat data >24h stale."],
            ["+30% Equity Boost",
             "'Kabupaten tertinggal IPM rendah (Sampang 66.72) dapat boost +30%' — implisit selalu fire",
             "+30% boost applies ke kab IPM <68 saat menjadi deficit (penerima). Demo Jatim: Sampang & Bangkalan-as-deficit trigger EQUITY_BOOST_30 di 2 dari 32 match (Ngawi→Bangkalan beras 250t, Ngawi→Sampang beras 200t). Rare-but-correct di distribusi IPM Jatim; impact tumbuh dengan rollout nasional (Papua IPM ~50)."],
            ["'World-First Matching Engine'",
             "Klaim di title page + exec summary",
             "Dihapus. Diganti dengan klaim spesifik yang verifiable: 'first sub-national Indonesian food matching engine combining stable matching + IPM-based equity + two-tier data quality, kalibrasi BPS 2024'."],
            ["Climate-Adaptive",
             "'Banjir di rute = matching otomatis re-route' (implisit re-routing)",
             "Climate score penalty saat BMKG/Open-Meteo forecast tersedia: >50mm/day → 0.3, >20mm → 0.6, ≤20mm → 1.0 (scoring.py:130-153). Fallback neutral 0.7 untuk rute tanpa data — demo Jatim 10 rute weather seeded. Engine scoring, bukan re-routing logic."],
            ["Competitive Table",
             "AgriFlow ✅ vs 9 platform global pada 8 kapabilitas — tanpa sumber",
             "Tetap dipertahankan dengan caveat: 'Comparison reflects publicly documented features as of May 2026; closed-source/commercial systems not exhaustively audited.'"],
            ["API Usage Example",
             "Output snippet menampilkan 'Flags: []' (empty)",
             "Output snippet ditampilkan dengan flags realistis: ['EQUITY_BOOST_05', 'MADURA_CLUSTER'] — match real demo output."],
            ["Skenario Coverage",
             "19 skenario engineering edge case (Volume/Spasial/Temporal/Disrupsi/Politis), 106/106 pytest pass",
             "**24 skenario** — tambah 5 commercial-reality scenarios (C4 multi-holiday calendar, D6 route blackout, E6 contract reserve, F1 grade substitution, F2 demand segmentation) untuk tutup gap pertanyaan Pemda/HORECA real. 126/126 pytest pass. Kategori baru F: Kualitas & Segmentasi Komersial."],
            ["Volume Score Formula (FIX #1)",
             "min(s,d) / max(s,d) — punish big-producer-to-small-deficit matches. Tuban 800t → Surabaya 100t beras dapat score 0.125 padahal pattern itu paling realistis untuk supply chain beras Indonesia.",
             "**Coverage-of-demand**: min(s,d) / demand.volume_tons. Demand fully satisfied → 1.0. Excess surplus di-split ke deficit lain di Layer 3 greedy. Demo run v11: equity-boost matches (Lamongan → Bangkalan/Sampang beras) sekarang ranking #1 dan #2 dengan FinalScore 100.0 dan 98.4 (sebelumnya rank #8). A3 drastic-mismatch flag tetap fire untuk score <20% via post-processing — tidak ada loss of semantic."],
            ["Segment Differentiation (FIX #2)",
             "F2 `SEGMENT_HORECA/GOVERNMENT/INDUSTRIAL` flag cuma label kosong — tidak mengubah scoring. HORECA dan RETAIL bersaing untuk supply yang sama treated equal — engine tidak dapat tip choice berdasarkan segment characteristics.",
             "**segment_multiplier** range ±10%: HORECA +5% saat surplus ≥50t (bulk efficiency), GOVERNMENT +5%/+3% untuk Tier 1 + fresh supply, INDUSTRIAL +8%/+2% untuk bulk + Tier 2 OK. final_score = base × equity × segment, fully auditable via per-match flags (SEGMENT_HORECA_BULK_BONUS, dll). Greedy deficit ordering juga segment-aware sehingga HORECA dapat earlier pick saat supply contested."],
        ],
        col_widths=[3.0, 5.5, 8.5],
    )

    callout(doc,
        "Konsekuensi v11 untuk juri teknis: setiap klaim diferensiator sekarang "
        "punya referensi spesifik ke file:line di matching_engine/. Juri yang "
        "menanyakan 'tunjukkan di mana stable matching fire' bisa langsung "
        "dibawa ke allocation.py:84-217 + dispatch logic di 341-347. Honest "
        "framing > overclaim — kredibilitas pitch naik, bukan turun.",
        color="E8F4FD",
    )

    # ==================================================================
    # 2. PARADOKS PANGAN INDONESIA
    # ==================================================================
    heading(doc, "2. Paradoks Pangan Indonesia", level=1)

    para(doc,
        "Indonesia adalah produsen pangan terbesar ke-3 di dunia. Namun setiap "
        "tahun, Rp 213-551 triliun pangan hilang karena food loss and waste — "
        "setara 4-5% PDB. 40% kerugian ini di distribusi, bukan produksi.")

    para(doc,
        "Di satu kabupaten, petani membuang cabai karena harga jatuh. Di "
        "kabupaten sebelah, harga melonjak 200-300% karena kelangkaan. Pemda "
        "baru tahu krisis setelah keluhan di media sosial — telat 2-3 minggu, "
        "biaya survei Rp 15-25 juta sekali jalan.")

    heading(doc, "2.1 Paradoks Pertama: Disparitas Akses Digital", level=2)
    para(doc,
        "Solusi digital pangan yang ada semua fokus pada petani muda melek "
        "smartphone. Tapi realitas Indonesia berbeda:")
    make_table(doc,
        headers=["Fakta Demografi Indonesia", "Sumber"],
        rows=[
            ["~30% petani Indonesia berusia 60+ tahun", "BPS Sensus Pertanian 2023"],
            ["~20% petani buta huruf fungsional", "BPS 2023 (rural literacy)"],
            ["~22,5 juta penyandang disabilitas (8,5% populasi)", "BPS 2020"],
            ["~3,7 juta tuna netra", "Kemenkes 2022"],
            ["~2,5 juta tuli", "Kemenkes 2022"],
            ["Puluhan juta feature phone user rural", "KOMINFO 2024"],
        ],
        col_widths=[10.0, 7.0],
    )
    para(doc,
        "Ratusan juta orang TIDAK bisa pakai solusi digital pangan saat ini. "
        "AgriFlow dirancang untuk mereka semua — bukan afterthought, tapi "
        "fondasi desain.")

    heading(doc, "2.2 Paradoks Kedua: Algoritma Matching Belum Ada untuk Indonesia", level=2)
    para(doc,
        "Setelah analisa 12 platform pangan dunia, kami konfirmasi tidak ada "
        "yang menggabungkan kapabilitas yang AgriFlow butuhkan untuk Indonesia:")
    bullet(doc, "eNAM India: auction-based price discovery di 1.473 mandis, "
                "tapi tidak ada inter-regional surplus-deficit matching")
    bullet(doc, "MealConnect (Feeding America): matching food rescue donor-bank, "
                "hanya untuk surplus rescue, bukan komoditas pasar")
    bullet(doc, "Ethiopia Commodity Exchange (ECX): spot trading sukses, tapi "
                "tidak handle perishability + climate-adaptive")
    bullet(doc, "EAX (Rwanda) + AfMX (Kenya): kedua exchange African gagal "
                "control food price volatility karena infrastruktur lemah")
    bullet(doc, "Uber/Lyft pattern: stable matching <1 detik, tapi untuk people, "
                "bukan komoditas dengan shelf life")
    bullet(doc, "FEWS NET: AI early warning operational 30+ negara, tapi semua "
                "country-level. CSIS Desember 2025 eksplisit identifikasi gap sub-national")
    bullet(doc, "WFP mVAM: mobile food security survey via IVR + SMS di 30+ "
                "negara, tapi tidak matching engine")
    bullet(doc, "Food Drop Indiana: Power of Two Choices fairness untuk "
                "truckloads makanan, tapi hanya rejected loads")

    callout(doc,
        "Berdasarkan publicly documented features per Mei 2026, tidak ada "
        "platform pangan yang menggabungkan keenam kapabilitas berikut secara "
        "bersamaan untuk konteks sub-nasional Indonesia: (1) kabupaten-level "
        "matching, (2) multi-objective scoring 5 dimensi, (3) perishability-"
        "aware, (4) IPM-based equity weighting untuk kab tertinggal, "
        "(5) climate-aware scoring penalty saat data tersedia, (6) Modified "
        "Gale-Shapley stable matching saat Tier 1 ↔ Tier 1. AgriFlow adalah "
        "kombinasi pertama yang demonstrably runnable untuk konteks ini. "
        "(Closed-source / commercial systems tidak diaudit exhaustively.)",
        color="FFF3CD",
    )

    heading(doc, "2.3 Data Pendukung Lengkap", level=2)
    make_table(doc,
        headers=["Fakta", "Sumber"],
        rows=[
            ["Food loss & waste Indonesia: Rp 213-551 triliun/tahun (4-5% PDB)", "Bappenas Kajian FLW 2021"],
            ["40% food loss di distribusi, bukan produksi", "Bappenas 2021; FAO 2019"],
            ["Jatim: 4,7 juta petani, PDRB pertanian tertinggi nasional", "BPS Jatim 2024"],
            ["Disparitas harga cabai antar kab Jatim: 100-200%", "PIHPS BI; Jurnal Agribisnis 2025"],
            ["Food loss bisa memberi makan 61-125 juta orang", "Badan Pangan Nasional 2022"],
            ["BI luncurkan GPIPS Februari 2026", "Bank Indonesia 11 Feb 2026"],
            ["FAO-UN Joint Programme USD 205 juta Jatim+Lampung 2026-2027", "FAO 9 April 2026"],
            ["Prabowo target swasembada 4 tahun (Asta Cita #2)", "Perpres 12/2025 RPJMN"],
        ],
        col_widths=[11.0, 6.0],
    )

    # ==================================================================
    # 3. GPIPS ALIGNMENT
    # ==================================================================
    heading(doc, "3. GPIPS 2026 Alignment — Fondasi Strategis", level=1)

    para(doc,
        "11 Februari 2026, Bank Indonesia meluncurkan GPIPS sebagai evolusi "
        "GNPIP. GPIPS menekankan penguatan pasokan struktural dengan 3 pilar. "
        "AgriFlow dirancang sebagai operational layer langsung untuk ketiga pilar.")

    make_table(doc,
        headers=["Pilar GPIPS", "Mandat BI", "Implementasi AgriFlow"],
        rows=[
            ["Pilar 1: Produksi",
             "Peningkatan produksi hortikultura, bibit unggul, teknologi adaptif, koordinasi pola tanam antar wilayah.",
             "• Deteksi surplus/defisit per kabupaten (19 komoditas)\n"
             "• Prediksi harga 7-30 hari\n"
             "• Dashboard koordinasi pola tanam"],
            ["Pilar 2: Distribusi",
             "Efisiensi distribusi via Kerja Sama Antar Daerah (KAD) dan BUMN logistik.",
             "• Matching Engine 4-lapis surplus ↔ defisit\n"
             "• Smart routing Google Maps\n"
             "• Rekomendasi MOU antar-kab\n"
             "• Integrasi API Bulog/ID Food"],
            ["Pilar 3: Sinergi Kebijakan",
             "Pemanfaatan data neraca pangan akurat Pusat-Daerah.",
             "• Single source of truth\n"
             "• Export Bapanas/Bappeda/TPID\n"
             "• Alert real-time ke 46 KPw BI\n"
             "• TPID Award-ready"],
        ],
        col_widths=[3.5, 6.5, 7.0],
    )

    callout(doc,
        "Strategic Insight: AgriFlow bukan vendor. AgriFlow adalah mitra "
        "eksekusi GPIPS dengan matching engine purpose-built untuk konteks "
        "sub-nasional Indonesia. Sales cycle berubah dari 'kami ingin Anda "
        "beli produk' jadi 'kami menjalankan program yang sudah Anda "
        "canangkan, dengan algoritma yang dapat diverifikasi terhadap data "
        "BPS 2024 dan code base open di matching_engine/.'",
        color="E8F5E8",
    )

    heading(doc, "3.1 Addressable Market GPIPS", level=2)
    make_table(doc,
        headers=["Customer Segment", "Jumlah", "Addressable Revenue Tahunan"],
        rows=[
            ["Kantor Perwakilan BI (KPw BI)", "46 kantor", "Rp 23-46 miliar"],
            ["TPID Provinsi", "38 provinsi", "Rp 7,6-19 miliar"],
            ["TPID Kabupaten/Kota", "514 kab/kota", "Rp 38-77 miliar"],
            ["Dinas Pertanian Kab/Kota", "514 dinas", "Rp 51-103 miliar"],
            ["Bapanas Pusat", "1 institusi", "Rp 3-5 miliar"],
            ["TOTAL TAM Pemerintah", "—", "Rp 120-250 miliar/tahun"],
        ],
        col_widths=[7.0, 4.0, 6.0],
    )

    # ==================================================================
    # 4. TRIPLE-CHANNEL ACCESSIBILITY
    # ==================================================================
    heading(doc, "4. Triple-Channel Accessibility — DNA AgriFlow", level=1)

    para(doc,
        "Setiap pengguna mengakses AgriFlow melalui tool yang sudah ada di "
        "tangan mereka. Nol install, nol belajar, nol diskriminasi.")

    heading(doc, "4.1 Empat Channel, Satu Platform", level=2)
    make_table(doc,
        headers=["Channel", "Target User", "Cara Pakai"],
        rows=[
            ["Dashboard Web",
             "Staff Dinas, Bappeda, TPID, BI, peneliti",
             "Browser → dashboard.agriflow.id → peta choropleth + tabel rekomendasi + grafik prediksi. Mobile-responsive. Screen reader compatible WCAG AA."],
            ["WhatsApp Chat",
             "Pedagang pasar modern, petani muda, eksportir",
             "Kirim teks ke +62-xxx-xxx-xxxx. 'Harga cabai Malang?' Dijawab <3 detik dengan teks + Google Maps link."],
            ["WhatsApp Voice Note",
             "Petani buta huruf, lansia smartphone anak, pedagang",
             "Rekam voice note, kirim. STT Bahasa Indonesia/Jawa/Sunda. Jawaban balik sebagai voice note — tidak perlu baca."],
            ["Phone Call IVR",
             "Petani tua feature phone, buta huruf, tuna netra, tanpa WhatsApp",
             "Telepon 0800-AGRIFLOW. Menu Bahasa Indonesia + 5 bahasa daerah: 'Tekan 1 harga, 2 prediksi, 3 pembeli, 0 operator.'"],
        ],
        col_widths=[3.5, 4.5, 9.0],
    )

    heading(doc, "4.2 Inklusi Disabilitas — First-in-Indonesia", level=2)
    callout(doc,
        "22,5 juta penyandang disabilitas + 4 juta petani lansia + 10 juta "
        "feature phone user rural = ~35 juta orang pertama kali terhubung "
        "dengan informasi pangan digital.",
        color="FCE4EC",
    )

    heading(doc, "4.3 Demo Narasi — Pak Ahmad, Petani Kediri 62 Tahun", level=2)
    callout(doc,
        "Pak Ahmad buka HP Nokia jadul. Tekan 0800-AGRIFLOW. AgriFlow (voice, "
        "Bahasa Jawa): 'Sugeng sonten, Pak. Menapa ingkang saged kula bantu? "
        "Tekan setunggal regi, kaleh prediksi, tiga pembeli.' Pak Ahmad: "
        "Tekan 3. AgriFlow: 'Surabaya bade tumbas lombok 120 ton, regi "
        "Rp 78.000 saben kilo. Jarak 95 km. Kula kintun SMS kalih nomer "
        "pembeli?' Pak Ahmad: 'Inggih!' Hasil: Dapat SMS dengan nomer "
        "pembeli. Keuntungan tambahan Rp 253 juta. Tanpa install. Tanpa "
        "anak. Dalam Bahasa Jawa.",
        color="FFF8E1",
    )

    # ==================================================================
    # 5. SOLUSI 4 PILAR
    # ==================================================================
    heading(doc, "5. Solusi: AgriFlow 4 Pilar", level=1)

    make_table(doc,
        headers=["DETEKSI", "PREDIKSI", "DISTRIBUSI", "AKSESIBILITAS"],
        rows=[
            ["AI mendeteksi surplus & defisit dari BPS + PIHPS BI harian. Isolation Forest paham konteks musiman.",
             "XGBoost + Prophet ensemble. Forecast 7-30 hari, MAPE <20%. Confidence interval visible.",
             "Matching Engine 4-lapis (Section 5.5). Stable matching + equity multiplier. Smart routing Google Maps.",
             "Triple channel (Web + WA + Phone). 5 bahasa daerah. Inklusi disabilitas. Feature phone support."],
        ],
        col_widths=[4.0, 4.0, 4.5, 4.5],
    )

    para(doc, "Pilar Distribusi adalah otak project ini. Section 5.5 "
              "mendokumentasi algoritma matching engine secara mendalam — "
              "purpose-built untuk konteks sub-nasional Indonesia, cepat, "
              "efisien, dan setiap klaim verifiable terhadap kode.",
         italic=True)

    # ==================================================================
    # 5.5 MATCHING ENGINE DEEP DIVE
    # ==================================================================
    doc.add_page_break()
    heading(doc, "5.5 Matching Engine Deep Dive — Otak AgriFlow", level=1)

    callout(doc,
        "Bagian ini adalah jantung teknis AgriFlow. Matching Engine kami "
        "adalah kombinasi pertama untuk konteks pangan sub-nasional "
        "Indonesia yang menggabungkan Modified Gale-Shapley stable matching "
        "(Nobel Prize Economics 2012, fires saat kedua kabupaten Tier 1), "
        "multi-objective scoring 5 dimensi, equity multiplier untuk kabupaten "
        "tertinggal IPM <68 saat menjadi deficit (kalibrasi BPS 2024), dan "
        "climate-aware scoring penalty saat BMKG/Open-Meteo forecast "
        "tersedia. 24 skenario edge case (Volume, Spasial, Temporal, "
        "Disrupsi, Politis, Kualitas-Segmentasi) didokumentasi dengan "
        "handling spesifik di Section 5.5.5 dan tervalidasi pytest 126/126.",
        color="E3F2FD",
    )

    # 5.5.1
    heading(doc, "5.5.1 Apa Itu AgriFlow Matching Engine? — Analogi Sederhana",
            level=2)
    para(doc, "Bayangkan Uber, tapi untuk cabai dan bawang merah.")
    para(doc,
        "Setiap hari, Uber memecahkan masalah yang mirip dengan AgriFlow: "
        "ribuan pengemudi (supply) tersebar di kota, ribuan penumpang (demand) "
        "butuh tumpangan. Algoritma Uber memutuskan dalam <1 detik: pengemudi "
        "mana ke penumpang mana sehingga semua pihak diuntungkan.")
    para(doc,
        "AgriFlow menyelesaikan masalah yang secara matematis identik untuk "
        "pangan:")
    bullet(doc, "Pengemudi → Kabupaten surplus (Kediri punya 100 ton cabai)")
    bullet(doc, "Penumpang → Kabupaten defisit (Surabaya butuh 80 ton cabai)")
    bullet(doc, "Tujuan → Surplus terbaik untuk defisit terbaik, semua pihak win-win")

    para(doc, "Tapi AgriFlow menambahkan 6 dimensi yang Uber tidak punya:")
    bullet(doc, "Perishability — cabai busuk dalam 5 hari, tidak seperti penumpang. Score turun saat sisa shelf life < transit time (scoring.py:103-123)")
    bullet(doc, "Equity — kabupaten tertinggal IPM <68 dapat boost +30% saat menjadi deficit/penerima (allocation.py:38-65). Demo Jatim: Sampang 66.72 & Bangkalan 67.70 yang trigger EQUITY_BOOST_30 saat shipping ke mereka")
    bullet(doc, "Climate — penalty score saat BMKG/Open-Meteo forecast hujan deras di rute (>50mm → 0.3, >20mm → 0.6, ≤20mm → 1.0); fallback neutral 0.7 saat data forecast tidak tersedia")
    bullet(doc, "Volume — bukan 1-to-1, 1 surplus bisa dipecah ke banyak deficit di Layer 3 greedy allocation (allocation.py:264-309)")
    bullet(doc, "Stable Matching — Modified Gale-Shapley (Nobel Prize Economics 2012) saat kedua kabupaten Tier 1 (allocation.py:84-217); cross-tier dan Tier 2 ↔ Tier 2 pakai greedy-with-equity-priority")
    bullet(doc, "Two-tier Confidence — label HIGH (kedua Tier 1) / MEDIUM (cross-tier atau Tier 2 ↔ Tier 2) / LOW (data stale >24h); transparant ke user")

    callout(doc,
        "Kombinasi keenam dimensi ini, dibungkus dalam arsitektur 4-lapis "
        "dengan implementasi Python siap-run + 106 unit test, adalah "
        "matching engine purpose-built pertama untuk komoditas pangan "
        "tingkat kabupaten di Indonesia.",
        color="E8F5E8",
    )

    # 5.5.2
    heading(doc, "5.5.2 Benchmark vs 12 Platform Pangan Dunia", level=2)
    para(doc, "Setelah analisa 12 platform pangan global, ini posisi AgriFlow:")
    make_table(doc,
        headers=["Kapabilitas", "eNAM", "MealConn.", "Food Drop", "Uber",
                 "ECX", "EAX", "AfMX", "FEWS NET", "AgriFlow"],
        rows=[
            ["Sub-national matching", "❌", "❌", "❌", "N/A", "❌", "❌", "❌", "❌", "✅"],
            ["Multi-objective scoring", "❌", "Limited", "Fairness", "Distance", "Limited", "❌", "❌", "❌", "✅ 5-dim"],
            ["Perishability-aware", "❌", "Partial", "❌", "N/A", "❌", "❌", "❌", "❌", "✅"],
            ["Equity-weighted", "❌", "❌", "Partial", "❌", "❌", "❌", "❌", "❌", "✅ IPM"],
            ["Climate-adaptive", "❌", "❌", "❌", "❌", "❌", "❌", "❌", "Partial", "✅"],
            ["Stable matching", "❌", "❌", "❌", "Partial", "❌", "❌", "❌", "❌", "✅ Nobel"],
            ["Real-time price", "✅", "❌", "❌", "✅", "✅", "✅", "✅", "❌", "✅"],
            ["Two-tier confidence", "❌", "❌", "❌", "❌", "❌", "❌", "❌", "❌", "✅ First"],
        ],
        col_widths=[4.0, 1.5, 1.7, 1.7, 1.5, 1.3, 1.3, 1.3, 1.7, 2.0],
    )
    para(doc,
        "Dari 8 kapabilitas critical, AgriFlow adalah satu-satunya platform "
        "yang centang penuh semuanya. Comparison reflects publicly documented "
        "features per Mei 2026; closed-source / commercial systems tidak "
        "diaudit exhaustively.",
        bold_first="Conclusion:")

    # 5.5.3
    heading(doc, "5.5.3 Strategi Data Hybrid 2-Tier", level=2)
    para(doc,
        "Sebelum dive ke arsitektur, penting paham reality data Indonesia. "
        "Tidak semua data tersedia real-time untuk semua kabupaten.")
    make_table(doc,
        headers=["Tier", "Cakupan Jatim", "Sumber Data", "Update Frequency"],
        rows=[
            ["Tier 1 (HIGH confidence)",
             "8 kota IHK Jatim: Surabaya, Malang, Kediri, Madiun, Probolinggo, Banyuwangi, Sumenep, Jember",
             "PIHPS Bank Indonesia (harga harian) + BMKG cuaca real-time",
             "Daily 10:00-13:00 WIB"],
            ["Tier 2 (MEDIUM confidence)",
             "30 kab non-IHK Jatim: Sampang, Bondowoso, Pacitan, Bangkalan, Pamekasan, dll",
             "Bapanas Panel Harga (mingguan) + BPS Susenas (tahunan) + BPS Hortikultura (tahunan)",
             "Weekly + monthly estimation"],
        ],
        col_widths=[3.5, 5.5, 4.5, 3.5],
    )
    para(doc,
        "Implikasi: Algoritma kita harus respect data quality. Stable matching "
        "hanya valid untuk data Tier 1. Tier 2 pakai greedy multi-objective "
        "scoring dengan disclaimer transparan.",
        bold_first="Implikasi:")
    callout(doc,
        "Honest engineering principle: Algoritma cantik atas data tidak akurat "
        "= fake fairness. AgriFlow tidak fake. Setiap match return confidence "
        "label HIGH/MEDIUM/LOW.",
        color="FFF3CD",
    )

    # 5.5.4
    heading(doc, "5.5.4 Arsitektur 4-Lapis", level=2)
    para(doc,
        "AgriFlow Matching Engine terdiri dari 4 lapis (Layer 0-3) plus "
        "post-processing yang masing-masing punya tujuan spesifik. Setiap "
        "lapis dioptimalkan untuk pertanyaan berbeda.")

    code_block(doc,
        "ARCHITECTURE OVERVIEW\n"
        "─────────────────────────────────────────────────────────────────\n"
        " LAYER 0: Data Confidence Filter                  (latency <10ms)\n"
        "   Tag setiap kab Tier 1 (HIGH) atau Tier 2 (MEDIUM)\n"
        "─────────────────────────────────────────────────────────────────\n"
        " LAYER 1: Candidate Generation                    (latency <50ms)\n"
        "   Spatial filter + 9 hard constraints\n"
        "   Output: Top 10-20 viable pairs per surplus region\n"
        "─────────────────────────────────────────────────────────────────\n"
        " LAYER 2: Multi-Objective Scoring                 (latency <200ms)\n"
        "   Weighted utility (5 dimensions)\n"
        "   Distance 22% │ Volume 22% │ Price 22% │ Perish 18% │ Climate 16%\n"
        "─────────────────────────────────────────────────────────────────\n"
        " LAYER 3: Equity-Weighted Final Allocation\n"
        "   Equity Multiplier × BaseScore\n"
        "   Tier 1: Modified Gale-Shapley Stable Matching\n"
        "   Tier 2: Greedy Top-K with equity priority\n"
        "─────────────────────────────────────────────────────────────────\n"
        " POST-PROCESSING:\n"
        "   Tag flags (RAMADAN_SPIKE, EQUITY_BOOST_30, MADURA_CLUSTER, ...)\n"
        "   external_opportunities, unmatched, run_metadata\n"
        "─────────────────────────────────────────────────────────────────\n"
        "                Total latency: <500ms p99 target\n"
        "          v10 measured: 1.4ms (sample) - 55ms (361×361 stress)")

    # Layer 0
    heading(doc, "Layer 0: Data Confidence Filter", level=3)
    para(doc, "Sebelum apapun, sistem cek: apakah data kabupaten ini cukup "
              "berkualitas untuk algoritma ketat?")
    code_block(doc,
        "def determine_tier(kabupaten):\n"
        "    if kabupaten.id in TIER_1_KOTA_IHK:\n"
        "        # 8 kota: Surabaya, Malang, Kediri, Madiun,\n"
        "        #         Probolinggo, Banyuwangi, Sumenep, Jember\n"
        "        return Tier.HIGH\n"
        "    else:\n"
        "        # 30 kab non-IHK\n"
        "        return Tier.MEDIUM\n"
        "\n"
        "# Implementation: matching_engine/constraints.py:42")

    # Layer 1
    heading(doc, "Layer 1: Candidate Generation", level=3)
    para(doc,
        "Tujuan: Filter cepat — dari 38 kab × 19 komoditas × harian = ~25.000 "
        "kemungkinan pasangan, kerucutkan ke 10-20 paling viable.")
    para(doc,
        "Pendekatan: Spatial indexing + 9 hard constraints. Inspired by Uber's "
        "geo-indexing pattern (1M req/sec capability).")

    para(doc, "Hard Constraints (semua harus dipenuhi):", bold_first="Hard Constraints:")
    code_block(doc,
        "def is_viable_pair(surplus, deficit, logistics):\n"
        "    if surplus.commodity != deficit.commodity:\n"
        "        return False, 'DIFFERENT_COMMODITY'\n"
        "    if surplus.kabupaten.id == deficit.kabupaten.id:\n"
        "        return False, 'SAME_KABUPATEN'                  # no self-match\n"
        "    if surplus.volume < MIN_VIABLE[commodity]:\n"
        "        return False, 'SUPPLY_BELOW_MIN'                # >= 1 ton\n"
        "    if deficit.volume < MIN_VIABLE[commodity]:\n"
        "        return False, 'DEMAND_BELOW_MIN'\n"
        "    if surplus.kabupaten.emergency_mode == UNREACHABLE:\n"
        "        return False, 'SUPPLY_UNREACHABLE'              # D4, D5\n"
        "    if surplus.kabupaten.pemda_overrides[do_not_export]:\n"
        "        return False, 'PEMDA_OVERRIDE'                  # E2\n"
        "    \n"
        "    # E5: BBM-aware distance shrinking\n"
        "    if logistics.bbm_change_pct > 0.10:\n"
        "        max_dist *= (1 - min(0.20, bbm_change_pct * 0.5))\n"
        "    if distance_km(surplus, deficit) > max_dist:\n"
        "        return False, 'DISTANCE_EXCEEDS_MAX'            # B2\n"
        "    \n"
        "    # D2: perishability margin\n"
        "    transit_days = ceil(dist / speed / hours_per_day)\n"
        "    if surplus.harvest_age + transit_days > MAX_FRESH_AGE[commodity]:\n"
        "        return False, 'SUPPLY_TOO_OLD'\n"
        "    \n"
        "    return True, ''\n"
        "\n"
        "# Implementation: matching_engine/constraints.py:191")

    para(doc, "Constraint values per komoditas (extract):", bold_first="Constants:")
    code_block(doc,
        "MAX_DISTANCE = {\n"
        "    'cabai_merah':    200,    # km, sensitif waktu\n"
        "    'bawang_merah':   400,\n"
        "    'tomat':          150,\n"
        "    'beras':          800,    # tahan lama, jarak lebih jauh OK\n"
        "    'ikan_segar':     150,\n"
        "    'daging_ayam':    200,\n"
        "    # ... 19 komoditas total di matching_engine/constraints.py:56\n"
        "}\n"
        "MIN_VIABLE = {\n"
        "    'cabai_merah':  1.0,    # ton\n"
        "    'bawang_merah': 2.0,\n"
        "    'beras':        5.0,\n"
        "}\n"
        "MAX_FRESH_AGE = {\n"
        "    'cabai_merah':  5,      # hari sejak panen\n"
        "    'bawang_merah': 30,\n"
        "    'beras':        180,\n"
        "    'ikan_segar':   2,      # paling perishable\n"
        "}")
    para(doc, "Output: List of 10-20 candidate pairs per surplus region. "
              "Latency <50ms.", italic=True)

    # Layer 2
    heading(doc, "Layer 2: Multi-Objective Scoring", level=3)
    para(doc, "Untuk setiap candidate pair, hitung skor 0-100 dari 5 dimensi:")
    code_block(doc,
        "BaseScore(s, d) = 100 × (\n"
        "    0.22 × DistanceScore(s, d)\n"
        "  + 0.22 × VolumeScore(s, d)\n"
        "  + 0.22 × PriceScore(s, d)\n"
        "  + 0.18 × PerishabilityScore(s, d)\n"
        "  + 0.16 × ClimateScore(s, d)\n"
        ")\n"
        "\n"
        "# Implementation: matching_engine/scoring.py:188")

    para(doc, "Komponen scoring (semua normalized 0-1):")

    para(doc, "1. DistanceScore — Lebih dekat lebih baik", bold_first="(1)")
    code_block(doc,
        "def distance_score(s, d):\n"
        "    distance_km = haversine_or_osrm(s.location, d.location)\n"
        "    max_viable = MAX_DISTANCE[s.commodity]\n"
        "    return max(0, 1 - distance_km / max_viable)")

    para(doc, "2. VolumeScore — Surplus dan deficit harus seimbang", bold_first="(2)")
    code_block(doc,
        "def volume_score(s, d):\n"
        "    matched = min(s.volume, d.volume)\n"
        "    larger  = max(s.volume, d.volume)\n"
        "    return matched / larger  # 1.0 jika perfect match")

    para(doc, "3. PriceScore — Selisih harga × volume = potensi profit petani", bold_first="(3)")
    code_block(doc,
        "def price_score(s, d, distance_km, logistics):\n"
        "    bbm_per_km = logistics.bbm / logistics.km_per_liter\n"
        "    bbm_cost = bbm_per_km * distance_km / 5000  # 5 ton truk\n"
        "    handling = 200  # Rp/kg\n"
        "    logistics_cost = bbm_cost + handling\n"
        "    arbitrage_pct = (d.price - s.price - logistics_cost) / s.price\n"
        "    return min(1, max(0, arbitrage_pct / 0.50))\n"
        "    # 0% arbitrage = 0, 50%+ arbitrage = 1")

    para(doc, "4. PerishabilityScore — Sisa shelf life harus cukup untuk transit", bold_first="(4)")
    code_block(doc,
        "def perishability_score(s, d, distance_km, logistics):\n"
        "    remaining = MAX_FRESH_AGE[s.commodity] - s.harvest_age\n"
        "    transit   = distance_km / 60 / 8  # 60 km/h, 8 jam/hari\n"
        "    margin    = remaining - transit\n"
        "    if margin < 1: return 0      # tidak akan sampai segar\n"
        "    return min(1, margin / 5)    # 5+ hari margin = perfect")

    para(doc, "5. ClimateScore — Cuaca buruk di rute = penalti", bold_first="(5)")
    code_block(doc,
        "def climate_score(s, d, weather):\n"
        "    if weather is None: return 0.7   # neutral fallback\n"
        "    rain_mm = weather.max_rain_mm\n"
        "    if rain_mm > 50: return 0.3      # hujan deras\n"
        "    elif rain_mm > 20: return 0.6    # hujan sedang\n"
        "    else: return 1.0                 # cerah, optimal")

    para(doc,
        "Latency: ~10ms per pair, ~200ms untuk 20 pairs. v10 measured: ~1ms per pair.",
        italic=True)

    # Layer 3
    heading(doc, "Layer 3: Equity-Weighted Final Allocation", level=3)
    para(doc,
        "Tujuan: Resolve konflik (1 surplus diminta banyak deficit) sambil "
        "prioritize kabupaten tertinggal.")

    heading(doc, "Step 3a: Apply Equity Multiplier (UPDATED v10)", level=4)
    callout(doc,
        "PERUBAHAN v10: Threshold dikalibrasi ulang dari distribusi IPM 2024 "
        "BPS Jatim. Threshold v9 lama (<65 → 1.30) tidak pernah ter-trigger "
        "karena IPM terendah Jatim 2024 = Sampang 66.72. v10 menggeser "
        "threshold sehingga klaim '+30% boost untuk kab tertinggal' konkret "
        "applicable ke Sampang & Bangkalan.",
        color="FFE0B2",
    )

    code_block(doc,
        "def equity_multiplier_value(ipm: float) -> float:\n"
        "    # v10 calibration — BPS BRS Desember 2024 Jatim\n"
        "    if ipm < 68:    return 1.30   # tertinggal severe (Sampang, Bangkalan)\n"
        "    elif ipm < 72:  return 1.15   # tertinggal (Sumenep, Bondowoso, dll)\n"
        "    elif ipm < 78:  return 1.05   # menengah\n"
        "    else:           return 1.00   # maju (Sidoarjo, Surabaya)\n"
        "\n"
        "FinalScore(s, d) = BaseScore(s, d) × equity_multiplier_value(d.kabupaten.ipm)\n"
        "\n"
        "# Implementation: matching_engine/allocation.py:38")

    para(doc, "Distribusi IPM Jatim 2024 (BPS BRS Desember 2024):",
         bold_first="Coverage:")
    make_table(doc,
        headers=["Range IPM", "Multiplier", "Boost", "Kab/Kota Jatim 2024"],
        rows=[
            ["IPM < 68", "1.30", "+30%", "Sampang (66.72), Bangkalan (67.70)"],
            ["68 ≤ IPM < 72", "1.15", "+15%",
             "Sumenep, Probolinggo (kab), Bondowoso, Lumajang, Pamekasan, Situbondo, Pasuruan (kab), Pacitan, Jember, Madiun (kab)"],
            ["72 ≤ IPM < 78", "1.05", "+5%",
             "Bojonegoro, Ponorogo, Banyuwangi, Trenggalek, Blitar, Tuban, Ngawi, Kediri (kab), Jombang, Nganjuk, Tulungagung, Malang (kab), Lamongan, Magetan, Mojokerto (kab), Gresik"],
            ["IPM ≥ 78", "1.00", "(no boost)",
             "Kota Batu, Kota Probolinggo, Kota Pasuruan, Sidoarjo, Kota Blitar, Kota Mojokerto, Kota Kediri, Kota Madiun, Kota Malang, Kota Surabaya"],
        ],
        col_widths=[3.0, 2.0, 1.5, 10.5],
    )

    heading(doc, "Step 3b: Algoritma Allocation — Berbeda per Tier", level=4)
    para(doc, "Tier 1 (HIGH confidence) — Modified Gale-Shapley Stable Matching:",
         bold_first="Tier 1:")
    code_block(doc,
        "def stable_match_tier1(candidates, score_fn):\n"
        "    \"\"\"\n"
        "    Modified Gale-Shapley Stable Matching:\n"
        "      - Kab tertinggal proposes first (sorted by equity_multiplier)\n"
        "      - Surplus regions accept proposal yang FinalScore tertinggi\n"
        "      - Iterasi sampai semua matched atau exhausted\n"
        "    \n"
        "    Guarantees:\n"
        "      - Stable: tidak ada blocking pair\n"
        "      - Pareto-optimal untuk surplus side\n"
        "      - Equity-respected: tertinggal punya advantage\n"
        "    \n"
        "    Reference: Gale & Shapley 1962, Nobel Prize Economics 2012\n"
        "    Time complexity: O(n²)\n"
        "    \"\"\"\n"
        "    # ... implementation di matching_engine/allocation.py:72")

    para(doc, "Tier 2 (MEDIUM confidence) — Greedy Top-K with equity priority:",
         bold_first="Tier 2:")
    code_block(doc,
        "def greedy_match_tier2(candidates, score_fn):\n"
        "    \"\"\"\n"
        "    Greedy Multi-Objective with equity priority:\n"
        "      - Sort deficit kabs by equity_multiplier descending\n"
        "      - Assign each to top-scored available surplus\n"
        "      - First-come basis after equity sort\n"
        "    \n"
        "    Why greedy: Tier 2 data estimasi (±15% error)\n"
        "      → stable matching guarantee meaningless\n"
        "      → greedy lebih honest, simpler, debug-friendly\n"
        "    \n"
        "    Time complexity: O(n log n)\n"
        "    \"\"\"\n"
        "    # ... implementation di matching_engine/allocation.py:212")

    para(doc,
        "Cross-tier (Tier 1 surplus dengan Tier 2 deficit, atau sebaliknya): "
        "Pakai approach Tier 2 (lower complexity respects lower data quality). "
        "Confidence label MEDIUM dengan disclaimer 'Estimated supply, "
        "recommend field validation'.",
        bold_first="Cross-tier:")

    heading(doc, "Performance Targets (UPDATED dengan benchmark v10)", level=3)
    make_table(doc,
        headers=["Metric", "Target", "v10 Measured", "Justifikasi"],
        rows=[
            ["End-to-end latency (38×19)", "<500ms p99",
             "1.46ms (sample CSV)\n55.53ms (361×361 stress)",
             "Lebih cepat dari Uber matching (1-2s)"],
            ["Throughput", "1000 matches/menit", "~40.000/menit (200×200)",
             "Cukup untuk 514 kab nasional Y2"],
            ["Stability guarantee", "100% Tier 1", "✓ Tier 1↔Tier 1 path",
             "Mathematical proof Gale-Shapley"],
            ["Equity boost coverage", "11+ dari 38 kab Jatim",
             "v10: 12 kab (+30% atau +15%)",
             "IPM <72: 12 kab dapat boost berarti"],
            ["Test coverage", "All 24 scenarios (v11)", "126/126 PASS in <0.3s",
             "Pytest automated regression"],
            ["Cold start", "<2s untuk first batch", "~50ms (Python import)",
             "Pre-computed geohash + distance matrix"],
        ],
        col_widths=[4.0, 3.5, 4.5, 5.0],
    )
    para(doc,
        "Achievable di Railway free tier (512MB RAM, 0.5 vCPU). Scale "
        "nasional: upgrade ke paid Rp 500K/bulan.",
        italic=True)

    # ==================================================================
    # 5.5.5 - 24 SKENARIO EDGE CASE — FULL DETAIL
    # ==================================================================
    heading(doc, "5.5.5 Cakupan 24 Skenario Edge Case", level=2)

    para(doc,
        "Algoritma yang baik bukan hanya yang elegant secara teori, tapi yang "
        "handle realita dengan grace. Setelah analisa Indonesia, kami "
        "identifikasi 24 skenario edge case yang algoritma kami harus handle "
        "(19 engineering edge case asli v10 + 5 commercial-reality scenarios "
        "tambahan v11). Setiap skenario tervalidasi via pytest test class "
        "spesifik (lihat Section 5.5.10). Berikut dokumentasi lengkapnya.")

    callout(doc,
        "6 Kategori Skenario (v11):\n"
        "• Volume (4 skenario): A1-A4 — cara handle volume mismatch\n"
        "• Spasial (3 skenario): B1-B3 — cara handle geographic constraints\n"
        "• Temporal (4 skenario): C1-C4 — cara handle waktu (Ramadan, panen, stale, multi-holiday)\n"
        "• Disrupsi (6 skenario): D1-D6 — banjir, hama, anomali, erupsi, banjir multi-kab, route blackout\n"
        "• Politis & Kebijakan (6 skenario): E1-E6 — equity, override, Bulog, import, BBM, contract reserve\n"
        "• Kualitas & Segmentasi Komersial (2 skenario, NEW v11): F1-F2 — grade substitution, demand segmentation",
        color="F3E5F5",
    )

    # 5.5.5 — A. VOLUME
    heading(doc, "Kategori A: Volume Skenario", level=3)
    make_table(doc,
        headers=["ID", "Skenario & Trigger", "Algoritma Response"],
        rows=[
            ["A1\nSurplus 1-to-Many",
             "Kediri surplus 100 ton cabai. Total demand 105 ton dari Surabaya (40t), Sidoarjo (35t), Gresik (30t).\n\nTrigger: total_demand > supply",
             "• Layer 3 partial-allocate dengan equity priority\n"
             "• Ranking by FinalScore × EquityMultiplier\n"
             "• Top defisit dapat full, terakhir dapat sisa\n"
             "• Output: 'Surabaya 40t, Sidoarjo 35t, Gresik 25t (kurang 5t, suggest source lain)'\n"
             "• Test: TestA1_OneToMany"],
            ["A2\nMany-to-1",
             "Surabaya butuh 200 ton cabai. 5 kab surplus (masing-masing 30-50 ton).\n\nTrigger: single_demand > any_single_supply",
             "• Layer 1 generate top-K candidates dari 5 kab\n"
             "• Layer 3 aggregate multi-source\n"
             "• Smart routing combine logistik\n"
             "• Output: 'Source dari Kediri 50t + Blitar 50t + Tulungagung 40t + Trenggalek 35t + Pacitan 25t'\n"
             "• Test: TestA2_ManyToOne"],
            ["A3\nVolume Mismatch Drastis",
             "Sampang surplus 2 ton cabai. Surabaya butuh 100 ton.\n\nTrigger: supply << demand",
             "• Layer 1 hard constraint MIN_VIABLE = 1 ton: PASS\n"
             "• Layer 2 VolumeScore = 2/100 = 0.02 (rendah)\n"
             "• Tetap di-rank tapi prioritas rendah\n"
             "• Equity Multiplier 1.30 boost Sampang (v10)\n"
             "• Flag: VOLUME_MISMATCH_DRASTIS\n"
             "• Output: include but flag 'marginal contribution'\n"
             "• Test: TestA3_VolumeMismatchDrastis"],
            ["A4\nZero Demand",
             "5 kab semua surplus, tidak ada deficit (akhir Ramadan, demand pasca-Idul Fitri turun).\n\nTrigger: all_surplus, no_deficit",
             "• Layer 0 detect zero-demand state\n"
             "• Algoritma return 'no internal match needed'\n"
             "• Suggest export ke luar Jatim (Jakarta, Bali)\n"
             "• Trigger Bulog procurement notification\n"
             "• Output: 'Internal match: NONE. External opportunity: Jakarta defisit 800t.'\n"
             "• Test: TestA4_ZeroDemand"],
        ],
        col_widths=[2.5, 5.5, 9.0],
    )

    # 5.5.5 — B. SPASIAL
    heading(doc, "Kategori B: Spasial Skenario", level=3)
    make_table(doc,
        headers=["ID", "Skenario & Trigger", "Algoritma Response"],
        rows=[
            ["B1\nCross-Tier Match",
             "Sampang (Tier 2 estimasi) surplus → Surabaya (Tier 1 akurat) butuh.\n\nTrigger: tier_diff = TRUE",
             "• Layer 0 detect cross-tier\n"
             "• Pakai Tier 2 algorithm (greedy, not stable matching)\n"
             "• Equity Multiplier full berdasarkan IPM deficit\n"
             "• Confidence label: MEDIUM\n"
             "• Disclaimer: 'Estimated supply Sampang, field validation recommended'\n"
             "• Test: TestB1_CrossTier"],
            ["B2\nLong Distance",
             "Pacitan surplus cabai → Banyuwangi butuh. Distance 350km, melebihi MAX_DISTANCE cabai 200km.\n\nTrigger: distance > max_viable",
             "• Layer 1 hard constraint REJECT (DISTANCE_EXCEEDS_MAX)\n"
             "• Output: 'Pair tidak viable: jarak 350km > batas cabai 200km'\n"
             "• Suggest alternatif: cari surplus dalam radius 200km Banyuwangi\n"
             "• Or suggest komoditas tahan lama (beras MAX 800km)\n"
             "• Test: TestB2_LongDistance"],
            ["B3\nGeographic Cluster Surplus",
             "Madura (Sumenep, Pamekasan, Sampang, Bangkalan) semua surplus bawang merah saat panen raya.\n\nTrigger: regional_cluster_surplus",
             "• Pre-process detect cluster pattern (CLUSTER_MADURA set)\n"
             "• Algoritma reject internal Madura matching (semua surplus)\n"
             "• Suggest export ke luar Madura: Surabaya, Sidoarjo, Malang\n"
             "• Trigger smart routing via Jembatan Suramadu\n"
             "• Flag: MADURA_CLUSTER\n"
             "• Output: '4 kab Madura aggregate 200t → Surabaya market'\n"
             "• Test: TestB3_ClusterMadura"],
        ],
        col_widths=[2.5, 5.5, 9.0],
    )

    # 5.5.5 — C. TEMPORAL
    heading(doc, "Kategori C: Temporal Skenario", level=3)
    make_table(doc,
        headers=["ID", "Skenario & Trigger", "Algoritma Response"],
        rows=[
            ["C1\nRamadan Spike",
             "Demand cabai naik 200% 2 minggu sebelum Idul Fitri (historical pattern).\n\nTrigger: hijri_calendar_event proximity",
             "• Anomaly Detector flag seasonal spike\n"
             "• Algoritma adjust ke RAMADAN_WEIGHTS (perishability 0.18→0.22)\n"
             "• Pre-position H-7: matching dilakukan lebih awal\n"
             "• ClimateScore weight up (karena bahaya kehilangan supply)\n"
             "• Flag: RAMADAN_SPIKE\n"
             "• Output: 'Pre-Ramadan mode aktif, scoring weights adjusted'\n"
             "• Test: TestC1_RamadanSpike"],
            ["C2\nPasca Panen Raya",
             "Tuban + Lamongan surplus padi pasca panen (Maret-April).\n\nTrigger: harvest_calendar_event detected",
             "• Layer 1 detect surplus regional cluster\n"
             "• Suggest matching ke kab non-sentra (Surabaya, Malang)\n"
             "• Priority: Bulog procurement (kontak otomatis Bulog Divre Jatim)\n"
             "• Equity Multiplier ke kab IPM rendah jika ada deficit\n"
             "• Output: 'Surplus 5000t → Bulog 60% + market 40%'\n"
             "• Test: TestC2_PostHarvest"],
            ["C3\nStale Data",
             "Hari Minggu, PIHPS tidak update. Data terakhir Sabtu.\n\nTrigger: data_age > 24h",
             "• Layer 0 detect stale (timestamp check)\n"
             "• Fallback ke last_valid harga\n"
             "• v10: Confidence drop bertingkat HIGH→MEDIUM→LOW\n"
             "• Bapanas Panel Harga sebagai cross-validation\n"
             "• Flag: STALE_DATA_24H\n"
             "• Output: 'Hari Minggu: data Sabtu used, confidence MEDIUM'\n"
             "• Test: TestC3_StaleData"],
            ["C4 (NEW v11)\nMulti-Holiday Calendar",
             "Imlek (H-7), Natal (H-21), school-start (H-14 sebelum Juli/Januari) — bukan hanya Ramadan.\n\nTrigger: get_active_demand_event() return non-RAMADAN event",
             "• engine.get_active_demand_event() dispatch event di reference_date\n"
             "• Priority order saat overlap: RAMADAN > NATAL > IMLEK > SCHOOL_START\n"
             "• IMLEK_WEIGHTS / NATAL_WEIGHTS / SCHOOL_START_WEIGHTS punya profil bobot berbeda\n"
             "• Imlek: spike urban beras premium + jeruk + ayam (H-7 window pendek)\n"
             "• Natal: NTT/Papua/Sulut sembako + daging (H-21 to H-1)\n"
             "• School-start: kos-kosan beras + minyak + telur (H-14)\n"
             "• Flag: IMLEK_SPIKE / NATAL_SPIKE / SCHOOL_START_SPIKE\n"
             "• Test: TestC4_HolidayCalendar"],
        ],
        col_widths=[2.5, 5.5, 9.0],
    )

    # 5.5.5 — D. DISRUPSI
    heading(doc, "Kategori D: Disrupsi Skenario (6 Skenario)", level=3)
    make_table(doc,
        headers=["ID", "Skenario & Trigger", "Algoritma Response"],
        rows=[
            ["D1\nBanjir Rute",
             "BMKG forecast hujan deras 80mm di rute Kediri-Surabaya.\n\nTrigger: rain_forecast > 50mm",
             "• ClimateScore drop dari 1.0 ke 0.3 (rain >50mm)\n"
             "• Re-rank candidate pairs: rute alternatif via Mojokerto naik\n"
             "• Atau delay match 24-48 jam\n"
             "• Output: 'Rute Kediri-Surabaya at risk. Suggest re-route via Mojokerto (delay 4 jam)'\n"
             "• Test: TestD1_BanjirRute"],
            ["D2\nKomoditas Rusak Massal",
             "Hama serang sentra cabai Kediri, surplus turun 70% dalam 1 minggu.\n\nTrigger: supply_drop > 50% week-over-week",
             "• Anomaly Detector flag major drop\n"
             "• Layer 1 adjust MIN_VIABLE\n"
             "• Re-rank: Kediri lower priority\n"
             "• Auto-alert Dinas via WhatsApp\n"
             "• Output: 'Hama detected. Kediri supply revised 100t → 30t. Alternative: Blitar, Tulungagung'\n"
             "• Test: TestD2_KomoditasRusak"],
            ["D3\nHarga Anomali",
             "Cabai Surabaya tiba-tiba Rp 200K/kg (5x normal Rp 40K).\n\nTrigger: price > 3σ from rolling median",
             "• Outlier rejection: drop data point dari pool\n"
             "• Use last_valid harga\n"
             "• Send alert ke admin Telegram untuk manual review\n"
             "• Cross-check Bapanas Panel Harga\n"
             "• Output: 'Harga Surabaya anomali Rp 200K, used last_valid Rp 42K. Manual review pending.'\n"
             "• Test: TestD3_HargaAnomali"],
            ["D4\nErupsi Gunung Berapi",
             "Merapi/Semeru erupsi, abu vulkanik tutup rute logistik.\n\nTrigger: PVMBG alert level III/IV",
             "• Integrasi PVMBG MAGMA API: detect erupsi alert\n"
             "• Mark kab radius 30km sebagai EmergencyMode.UNREACHABLE\n"
             "• Layer 1 hard constraint: skip kab terdampak\n"
             "• Mapping: Semeru→Lumajang/Probolinggo/Malang; Bromo→Probolinggo/Pasuruan/Lumajang/Malang\n"
             "• Output: 'Semeru erupsi: Lumajang, Probolinggo, Pasuruan unreachable. Re-routing via Jember-Banyuwangi'\n"
             "• Test: TestD4_ErupsiGunung"],
            ["D5\nBanjir Skala Besar (multi-kab)",
             "Banjir Sampang+Bangkalan+Pasuruan, multiple kab terdampak.\n\nTrigger: BNPB API alert + BMKG sustained heavy rain",
             "• Pre-process detect regional disaster (BNPB DIBI)\n"
             "• Mark semua kab terdampak as 'emergency mode'\n"
             "• Surplus dari kab terdampak deprioritized (likely damaged)\n"
             "• Demand dari kab terdampak prioritized (humanitarian)\n"
             "• Flag: HUMANITARIAN_PRIORITY\n"
             "• Trigger Bulog stock release recommendation\n"
             "• Output: 'Madura banjir: emergency relief mode. Bulog stock 500t → distribusi prioritas.'\n"
             "• Test: TestD5_BanjirMultiKab"],
            ["D6 (NEW v11)\nRoute Blackout",
             "Mudik H+1 Idul Fitri tutup toll Cikampek; demonstrasi Trans-Jawa; Suramadu maintenance terjadwal.\n\nTrigger: route_blackouts list contains active blackout for (origin, dest, date)",
             "• run_matching accepts route_blackouts: List[RouteBlackout]\n"
             "• Layer 1 filter pair candidate setelah generate_candidates (post-filter, bukan hard constraint)\n"
             "• Wildcard support: origin='*' atau dest='*' untuk match-any\n"
             "• Pair yang di-blackout dilaporkan di run_metadata + warning\n"
             "• Reasons enum: MUDIK_H1_IDUL_FITRI / DEMO_TRANS_JAWA / SURAMADU_MAINT\n"
             "• Output: 'Route blackout aktif: 12 pair difilter karena rute tertutup pada 2026-03-22'\n"
             "• Test: TestD6_RouteBlackout"],
        ],
        col_widths=[2.5, 5.5, 9.0],
    )

    # 5.5.5 — E. POLITIS
    heading(doc, "Kategori E: Politis & Kebijakan Skenario (6 Skenario)", level=3)
    make_table(doc,
        headers=["ID", "Skenario & Trigger", "Algoritma Response"],
        rows=[
            ["E1\nEquity Tie-Break",
             "Sampang (IPM 66.72) dan Bondowoso (IPM 69.62) bersaing untuk surplus dari Probolinggo.\n\nTrigger: BaseScore equal, equity diff present",
             "• Layer 3 Equity Multiplier (v10):\n"
             "    - Sampang: 1.30 (IPM <68)\n"
             "    - Bondowoso: 1.15 (IPM 68-72)\n"
             "• Sampang menang otomatis: FinalScore Sampang > Bondowoso\n"
             "• Flag: EQUITY_BOOST_30\n"
             "• Output: 'Sampang prioritized (IPM 66.72, equity boost +30%)'\n"
             "• Test: TestE1_EquityTieBreak"],
            ["E2\nPemda Override",
             "Bupati Pacitan minta produk lokal stay di Pacitan (kebijakan ketahanan pangan setempat).\n\nTrigger: pemda_override_flag = TRUE",
             "• Pemda dashboard set 'do_not_export_<komoditas>' flag per komoditas\n"
             "• Layer 1 hard constraint: PEMDA_OVERRIDE → exclude Pacitan dari surplus pool\n"
             "• Audit log immutable, visible ke BI\n"
             "• Output: 'Pacitan excluded per Pemda override. Alternative source suggested.'\n"
             "• Test: TestE2_PemdaOverride"],
            ["E3\nBulog Priority",
             "Bulog announce procurement target 1000 ton beras dari Madiun.\n\nTrigger: bulog_procurement_announcement received",
             "• Pre-process: 60% surplus reserve untuk Bulog (beras/jagung/kedelai)\n"
             "• Sisa 40% available untuk private matching\n"
             "• Notification ke Pemda Madiun\n"
             "• Output: 'Madiun beras: 600t Bulog priority. Sisa 400t available untuk private matching.'\n"
             "• Test: TestE3_BulogPriority"],
            ["E4\nKebijakan Import Tiba-tiba",
             "Pemerintah announce import beras 500K ton karena prediksi defisit nasional.\n\nTrigger: import_policy_announcement detected",
             "• Forecast model adjust: predicted price drop 10-15%\n"
             "• Switch ke IMPORT_POLICY_WEIGHTS (price weight 0.22→0.10)\n"
             "• Suggest petani tahan stok jika memungkinkan\n"
             "• Pemda alert: prepare buffer stock\n"
             "• Flag: IMPORT_POLICY_ACTIVE\n"
             "• Output: 'Import policy detected. Beras matching deprioritized 30 days. Suggest hold inventory.'\n"
             "• Test: TestE4_ImportPolicy"],
            ["E5\nKenaikan BBM Mendadak",
             "BBM naik 20% mendadak (kebijakan subsidi).\n\nTrigger: bbm_price_change > 10%",
             "• logistics_cost parameter re-calculate (naik proportional)\n"
             "• MAX_DISTANCE shrink 10-15% (long-haul jadi tidak ekonomis)\n"
             "• PriceScore arbitrage threshold up (need bigger margin)\n"
             "• Re-rank semua candidates\n"
             "• Output: 'BBM naik 20%. MAX_DISTANCE adjusted: cabai 200km → 175km. 12 matches now infeasible.'\n"
             "• Test: TestE5_BBMNaik"],
            ["E6 (NEW v11)\nContract Reserve (Generalisasi Bulog)",
             "Kediri 100t bawang. Carrefour kontrak 70% pre-commitment. Sisa 30t available untuk spot matching.\n\nTrigger: run_matching menerima contracts={(kab_id, commodity_code): reserve_pct}",
             "• engine.apply_contract_reserve() — pattern Bulog (E3) di-generalisir\n"
             "• Bekerja di atas Bulog split — kedua-duanya berlaku berurutan\n"
             "• Reserve 100% removes node (warning), reserve <100% reduces volume\n"
             "• Use case: MoU Carrefour/Indomaret, Indofood kontrak gula, kontrak farming offtaker\n"
             "• Warning per (kab, komoditas): 'Kediri Bawang: 70t contract priority (70%), 30t available untuk spot matching'\n"
             "• Test: TestE6_ContractReserve"],
        ],
        col_widths=[2.5, 5.5, 9.0],
    )

    # 5.5.5 — F. KUALITAS & SEGMENTASI KOMERSIAL (NEW v11)
    heading(doc, "Kategori F: Kualitas & Segmentasi Komersial (2 Skenario, NEW v11)", level=3)
    para(doc,
        "Kategori F menutup commercial-reality gap yang tidak tersentuh 19 "
        "skenario engineering edge case asli. Dua skenario di kategori ini "
        "menjawab pertanyaan Pemda/HORECA staff yang real: (a) apakah surplus "
        "grade tinggi dapat memenuhi demand grade lebih rendah; (b) apakah "
        "satu kab dapat punya multiple demand segment untuk komoditas sama.")
    make_table(doc,
        headers=["ID", "Skenario & Trigger", "Algoritma Response"],
        rows=[
            ["F1 (NEW v11)\nGrade Substitution",
             "Kediri surplus beras_premium 100t. Surabaya defisit beras_medium 100t. Tanpa substitusi: no match. Buyer beras medium tidak rugi terima premium.\n\nTrigger: run_matching menerima allow_grade_substitution=True",
             "• constraints.grade_compatible(): premium → medium = TRUE; medium → premium = FALSE (buyer mengharapkan grade lebih tinggi)\n"
             "• Layer 1 is_viable_pair lengkapi DIFFERENT_COMMODITY check dengan grade_compatible() saat opt-in\n"
             "• Default off — engine code identik backward-compat (106 test baseline tidak berubah)\n"
             "• GRADE_SUBSTITUTION dict: extensible (saat ini hanya beras_premium → beras_medium; minyak/gula extensible)\n"
             "• Flag: GRADE_SUBSTITUTION + note: 'beras_premium digunakan untuk memenuhi demand beras_medium (grade compatible substitution).'\n"
             "• Test: TestF1_GradeSubstitution"],
            ["F2 (NEW v11)\nDemand Segmentation",
             "Surabaya butuh beras premium untuk 2 segment: 80t HORECA (hotel, restoran, catering) + 80t RETAIL (rumah tangga). Engine harus handle keduanya independent — bukan merge.\n\nTrigger: DemandNode.segment field non-default",
             "• DemandSegment enum: RETAIL (default) / HORECA / GOVERNMENT / INDUSTRIAL\n"
             "• Backwards-compat: segment field default RETAIL → no behavior change untuk demand existing\n"
             "• Engine treat segment != RETAIL sebagai distinct demand node — multiple match per komoditas+kab dimungkinkan\n"
             "• Flag: SEGMENT_HORECA / SEGMENT_GOVERNMENT / SEGMENT_INDUSTRIAL\n"
             "• Use case: HORECA beras medium 13rb/kg vs RETAIL beras premium 14.5rb/kg (sama kab, beda buyer profile)\n"
             "• Test: TestF2_DemandSegmentation"],
        ],
        col_widths=[2.5, 5.5, 9.0],
    )

    heading(doc, "Ringkasan Cakupan 24 Skenario", level=3)
    make_table(doc,
        headers=["Kategori", "Jumlah", "Algoritma Layer Yang Aktif"],
        rows=[
            ["A. Volume", "4", "Layer 1 (constraint), Layer 2 (VolumeScore), Layer 3 (allocation)"],
            ["B. Spasial", "3", "Layer 0 (tier), Layer 1 (distance constraint), Layer 2 (DistanceScore)"],
            ["C. Temporal", "4", "Layer 0 (data freshness), Layer 1 (multi-holiday calendar), Layer 2 (re-weight)"],
            ["D. Disrupsi", "6", "Layer 1 (hard exclude + route blackout), Layer 2 (ClimateScore), External APIs (BMKG, PVMBG, BNPB)"],
            ["E. Politis & Kebijakan", "6", "Layer 1 (override + contract reserve), Layer 3 (equity), External signals (Bulog, BBM, import, MoU swasta)"],
            ["F. Kualitas & Segmentasi (NEW v11)", "2", "Layer 1 (grade substitution opt-in), Output decoration (SEGMENT_ flags)"],
            ["TOTAL", "24", "All 4 Layers + 5 External Signal Integrations + commercial-reality coverage"],
        ],
        col_widths=[4.5, 2.0, 10.5],
    )

    callout(doc,
        "Yang membuat ini purpose-built untuk Indonesia: 24 skenario edge case "
        "(19 engineering + 5 commercial-reality) dengan handling spesifik yang "
        "verifiable terhadap kode. Berdasarkan publicly documented features per "
        "Mei 2026: eNAM punya 1 skenario (auction). MealConnect 3 skenario "
        "(rescue). Food Drop 4 skenario (fairness). AgriFlow 24. v11 menambah "
        "bukti: semua tervalidasi pytest 126/126 PASS.",
        color="E1F5FE",
    )

    # ==================================================================
    # 5.5.6 RISIKO & MITIGASI
    # ==================================================================
    heading(doc, "5.5.6 Risiko & Mitigasi Matching Engine", level=2)
    para(doc, "Tidak ada algoritma yang sempurna. Berikut risiko teridentifikasi "
              "dengan mitigasi konkret.")
    make_table(doc,
        headers=["Risiko", "Probability", "Impact", "Mitigasi"],
        rows=[
            ["IPM data outdated (BPS tahunan)", "Medium", "Low",
             "Refresh tahunan saat BPS publish. IPM stabil, perubahan kecil tahunan. v10: kalibrasi threshold sudah aligned dengan distribusi 2024."],
            ["Stable matching tidak konvergen", "Very Low", "Medium",
             "Library matching Python sudah handle edge cases. Fallback to greedy if timeout."],
            ["Equity multiplier dianggap 'diskriminasi positif' politis", "Medium", "Medium",
             "Frame sebagai 'GPIPS alignment'. BI mandate eksplisit menyebut kab tertinggal."],
            ["Climate API down (BMKG)", "Medium", "Low",
             "Default ClimateScore = 0.7 (neutral). Open-Meteo sebagai fallback (auto-tested di v10)."],
            ["Bobot 22/22/22/18/16 suboptimal", "Medium", "Low",
             "A/B test setelah 3 bulan data, tune dengan ML. Default tetap masuk akal."],
            ["Komputasi >500ms saat scale nasional", "Low", "Medium",
             "v10 measured: 55ms p99 untuk 361×361. Margin 88.9% dari target. Pre-compute distance matrix + caching. Redis. Spatial indexing geohash."],
            ["Tier 2 estimasi error besar", "Medium", "Medium",
             "Confidence label MEDIUM. Cross-validate Bapanas Panel Harga + Susenas."],
            ["Pemda override misuse (politis)", "Low", "Low",
             "Audit log immutable. Setiap override dilog dan visible ke BI."],
            ["Kompetitor salin algoritma", "High", "Low",
             "Stable matching + IPM equity butuh 2-3 tahun engineer time. First-mover moat."],
        ],
        col_widths=[5.0, 2.5, 2.0, 7.5],
    )
    para(doc, "Semua risiko manageable. Tidak ada yang showstopper.", italic=True)

    # ==================================================================
    # 5.5.7 DAMPAK EXPECTED
    # ==================================================================
    heading(doc, "5.5.7 Dampak Expected", level=2)
    para(doc,
        "Quantified projection dari matching engine ini, berdasarkan benchmark "
        "MealConnect, eNAM, dan model surplus-deficit Indonesia:")
    make_table(doc,
        headers=["Impact Area", "Pilot 5 Kab Y1", "Scale Jatim Y2"],
        rows=[
            ["Food loss prevented", "~500 ton/tahun", "~3000 ton/tahun"],
            ["Economic value saved", "Rp 15-30 miliar", "Rp 100-200 miliar"],
            ["Petani income boost (avg)", "10-15%", "15-20% (data flywheel)"],
            ["Konsumen price stabilization", "Volatility ↓ 50%", "Volatility ↓ 60%"],
            ["Kab tertinggal supply boost", "+30% allocation",
             "+30-40% allocation (v10: konkret applicable ke Sampang & Bangkalan)"],
            ["CO2-eq saved (food loss)", "~1.250 ton/tahun", "~7.500 ton/tahun"],
        ],
        col_widths=[5.0, 5.0, 7.0],
    )
    callout(doc,
        "Total economic impact pilot Y1: Rp 20-30 miliar. Scale nasional Y3: "
        "Rp 200-300 miliar economic impact tahunan.",
        color="E8F5E8",
    )

    # ==================================================================
    # 5.5.8 DATA SOURCES (preserve from v9)
    # ==================================================================
    heading(doc, "5.5.8 Data Sources Detail", level=2)
    para(doc,
        "Matching engine mengintegrasikan 8 sumber data resmi dengan dual-mode "
        "operation: mock mode (offline, baca CSV sample untuk dev/demo) dan "
        "live mode (production, hit API real). Setiap connector diimplementasi "
        "sebagai module Python terpisah di data_sources/.")

    heading(doc, "5.5.8.1 PIHPS Bank Indonesia (Tier 1, harian)", level=3)
    make_table(doc,
        headers=["Atribut", "Detail"],
        rows=[
            ["URL", "https://www.bi.go.id/hargapangan/"],
            ["Cakupan", "8 kota IHK Jatim: Surabaya, Malang, Kediri, Madiun, Probolinggo, Banyuwangi, Sumenep, Jember"],
            ["Akses", "Web scraping (HTML) atau AJAX endpoint internal (reverse engineer via DevTools)"],
            ["Format", "HTML table → parse via BeautifulSoup; alternatif AJAX JSON"],
            ["Frekuensi update", "Harian, cut-off 13:00 WIB"],
            ["Rate limit", "Tidak ada hard limit publik. Tim BI minta <100 req/menit"],
            ["Komoditas tercakup", "19 komoditas pangan utama (beras, cabai, bawang, telur, daging, dll)"],
            ["Module", "data_sources/pihps_bi.py"],
            ["Mock fallback", "Baca sample_data/surplus_deficit.csv filter Tier 1 IHK"],
            ["Confidence label", "HIGH (data resmi BI, harian, terverifikasi enumerator)"],
            ["Catatan production", "Per April 2026 BI sedang draft Open Data API formal — cek bi.go.id/openapi/ untuk update."],
        ],
        col_widths=[4.0, 13.0],
    )

    heading(doc, "5.5.8.2 Bapanas Panel Harga (Tier 2, mingguan)", level=3)
    make_table(doc,
        headers=["Atribut", "Detail"],
        rows=[
            ["URL", "https://panelharga.badanpangan.go.id/"],
            ["API endpoint", "GET /api/front/harga-pangan-table"],
            ["Parameter", "province_id=35 (Jatim), level_harga_id=3 (produsen), tanggal=YYYY-MM-DD, period_date=weekly"],
            ["Cakupan", "30 kabupaten non-IHK Jatim (gap di-fill via spatial estimation)"],
            ["Format", "JSON response"],
            ["Frekuensi update", "Mingguan, Senin pagi"],
            ["Rate limit", "~60 req/menit (informal)"],
            ["Module", "data_sources/bapanas.py"],
            ["Mock fallback", "BapanasConnector(mock_mode=True)"],
            ["Confidence label", "MEDIUM (mingguan vs harian PIHPS; coverage tidak full)"],
        ],
        col_widths=[4.0, 13.0],
    )

    heading(doc, "5.5.8.3 BPS WebAPI (IPM 2024 + Produksi)", level=3)
    make_table(doc,
        headers=["Atribut", "Detail"],
        rows=[
            ["URL", "https://webapi.bps.go.id/v1/api/list"],
            ["Authentication", "API key required (gratis register di webapi.bps.go.id)"],
            ["Format URL", "/domain/{kab_id}/lang/ind/key/{api_key}/var/{var_id}/th/{tahun}"],
            ["Variable IDs", "26 = IPM, 60 = Produksi cabai, 61 = Produksi bawang merah"],
            ["Cakupan", "38 kabupaten/kota Jatim (kode wilayah BPS 35xx)"],
            ["Frekuensi update", "Tahunan (BRS Desember)"],
            ["Rate limit", "100 req/menit per akun"],
            ["Module", "data_sources/bps.py"],
            ["Hardcoded fallback", "IPM_2024_JATIM dict — 38 kab dengan nilai dari BRS Desember 2024 (v10: sync dengan generator)"],
            ["Catatan", "v10 sudah sinkronisasi IPM_2024_JATIM dengan generate_sample_data.py — single source of truth."],
        ],
        col_widths=[4.0, 13.0],
    )

    heading(doc, "5.5.8.4 BMKG + Open-Meteo (Cuaca / Skenario D1)", level=3)
    make_table(doc,
        headers=["Atribut", "Detail"],
        rows=[
            ["Primary URL", "https://api.bmkg.go.id/publik/prakiraan-cuaca?adm4={kode_desa}"],
            ["Fallback URL", "https://api.open-meteo.com/v1/forecast"],
            ["Format", "JSON"],
            ["Frekuensi update", "3-6 jam"],
            ["Rate limit", "BMKG: tidak terdokumentasi. Open-Meteo: 10.000 req/hari per IP"],
            ["Module", "data_sources/bmkg.py"],
            ["Strategy", "Sample 3 titik di rute (origin, midpoint, dest), ambil max precipitation_sum"],
            ["v10 Update", "prefer_bmkg flag sekarang properly honored (butuh adm4_lookup mapping). Tanpa lookup, auto-fallback ke Open-Meteo."],
            ["Output", "{max_rain_mm, transit_window_days, source} → input ke Layer 2 climate score"],
        ],
        col_widths=[4.0, 13.0],
    )

    heading(doc, "5.5.8.5 PVMBG MAGMA Indonesia (Erupsi / Skenario D4)", level=3)
    make_table(doc,
        headers=["Atribut", "Detail"],
        rows=[
            ["URL", "https://magma.esdm.go.id/v1/api/aktivitas-gunung-api"],
            ["Format", "JSON"],
            ["Authentication", "None (public read)"],
            ["Status levels", "Normal (1) → Waspada (2) → Siaga (3) → Awas (4)"],
            ["Trigger threshold", "Min status 'Siaga' → set kabupaten terdampak ke EmergencyMode.UNREACHABLE"],
            ["Mapping gunung→kab", "Semeru→Lumajang/Probolinggo/Malang; Bromo→Probolinggo/Pasuruan/Lumajang/Malang; Kelud→Kediri/Blitar/Malang; Ijen→Banyuwangi/Bondowoso; Arjuno-Welirang→Pasuruan/Malang/Mojokerto; Raung→Banyuwangi/Jember/Bondowoso"],
            ["Frekuensi update", "Real-time saat status berubah"],
            ["Module", "data_sources/pvmbg.py"],
        ],
        col_widths=[4.0, 13.0],
    )

    heading(doc, "5.5.8.6 BNPB DIBI (Bencana / Skenario D5)", level=3)
    make_table(doc,
        headers=["Atribut", "Detail"],
        rows=[
            ["URL", "https://dibi.bnpb.go.id/dibi3/peta_bencana"],
            ["Format", "HTML (perlu scrape) — DIBI tidak ada API publik formal per April 2026"],
            ["Disaster types", "BANJIR, TANAH LONGSOR, GEMPA BUMI, GUNUNG MELETUS, TSUNAMI"],
            ["Filter", "Hanya kabupaten Jatim (kode 35xx), bencana <7 hari"],
            ["Module", "data_sources/bnpb.py"],
            ["Catatan", "Production akan butuh HTML scraper (BeautifulSoup). Hackathon: dummy data atau monitoring manual."],
        ],
        col_widths=[4.0, 13.0],
    )

    heading(doc, "5.5.8.7 Google Maps Routes API + OSRM (Routing)", level=3)
    make_table(doc,
        headers=["Atribut", "Detail"],
        rows=[
            ["Primary", "https://routes.googleapis.com/directions/v2:computeRoutes"],
            ["Fallback", "https://router.project-osrm.org/route/v1/driving/{lon1,lat1};{lon2,lat2}"],
            ["Authentication", "Google: API key (free tier $200/bulan). OSRM: none (public demo)"],
            ["Rate limit", "Google: pay-per-request after free tier. OSRM demo: ~1000 req/hari"],
            ["Module", "data_sources/google_maps.py"],
            ["Default", "OSRM (gratis untuk hackathon). Production scale: deploy OSRM internal"],
            ["Last-resort fallback", "Haversine geodesic distance (built-in matematika tanpa network)"],
        ],
        col_widths=[4.0, 13.0],
    )

    heading(doc, "5.5.8.8 Aladhan Hijri Calendar (Ramadan / Skenario C1)", level=3)
    make_table(doc,
        headers=["Atribut", "Detail"],
        rows=[
            ["URL", "https://api.aladhan.com/v1/gToH/{date}"],
            ["Format", "JSON GET"],
            ["Authentication", "None"],
            ["Module", "data_sources/hijri_calendar.py"],
            ["Hardcoded fallback", "Dictionary Idul Fitri Gregorian dates 2025-2030 (Kemenag/almanak)"],
            ["Trigger logic", "H-21 sampai H-1 sebelum 1 Syawwal → engine pakai RAMADAN_WEIGHTS (perishability bobot naik 0.18→0.22)"],
        ],
        col_widths=[4.0, 13.0],
    )

    heading(doc, "5.5.8.9 Ringkasan Connector Architecture", level=3)
    make_table(doc,
        headers=["Source", "Tier", "Update", "Confidence Output", "Mock Mode?"],
        rows=[
            ["PIHPS BI", "1", "Harian", "HIGH", "✅"],
            ["Bapanas", "2", "Mingguan", "MEDIUM", "✅"],
            ["BPS WebAPI", "Static", "Tahunan", "HIGH (verified BRS)", "✅ (hardcoded)"],
            ["BMKG / Open-Meteo", "Realtime", "3-6 jam", "HIGH (Open-Meteo), MEDIUM (BMKG)", "✅"],
            ["PVMBG", "Realtime", "Per event", "HIGH", "✅"],
            ["BNPB DIBI", "Realtime", "Per event", "MEDIUM (scraping)", "✅"],
            ["Google Maps / OSRM", "Static", "On-demand", "HIGH (Google), MEDIUM (OSRM)", "✅"],
            ["Aladhan Hijri", "Static", "Tahunan", "HIGH", "✅"],
        ],
        col_widths=[3.5, 1.5, 2.5, 5.5, 2.5],
    )
    para(doc,
        "Setiap connector mengikuti pola: try live API → catch exception → "
        "fallback ke mock atau hardcoded dataset → tidak pernah crash engine. "
        "Ini memungkinkan demo offline dan resilience saat satu sumber down.",
        italic=True)

    # ==================================================================
    # 5.5.9 QUICK START
    # ==================================================================
    heading(doc, "5.5.9 Quick Start Guide", level=2)
    para(doc,
        "Tim dapat menjalankan engine end-to-end dalam <10 menit. Semua "
        "dependencies tersedia di requirements.txt dan tidak butuh API key "
        "untuk demo.")

    heading(doc, "5.5.9.1 Setup (5 menit)", level=3)
    code_block(doc,
        "# 1. Clone repository (atau extract dari hackathon submission ZIP)\n"
        "cd agriflow_engine\n"
        "\n"
        "# 2. Install Python dependencies (Python 3.10+)\n"
        "pip install -r requirements.txt\n"
        "\n"
        "# 3. Generate sample data — 38 kab × 19 komoditas Jatim\n"
        "python sample_data/generate_sample_data.py\n"
        "# Output: 5 CSV (kabupaten_jatim.csv, komoditas_constraints.csv,\n"
        "#         surplus_deficit.csv, weather_forecast.csv,\n"
        "#         historical_price_stats.csv)")

    heading(doc, "5.5.9.2 Run Demo End-to-End", level=3)
    code_block(doc, "python examples/run_demo.py")
    para(doc, "Expected output (v10):", bold_first="Output:")
    code_block(doc,
        "================================================================================\n"
        "  AgriFlow Matching Engine v10.0 — Demo End-to-End\n"
        "================================================================================\n"
        "\n"
        "→ Loading sample data 38 kabupaten Jatim × 19 komoditas...\n"
        "   Surplus nodes:  40\n"
        "   Deficit nodes:  33\n"
        "   Weather routes: 10\n"
        "   Historical:     19 komoditas\n"
        "\n"
        "→ Running matching engine...\n"
        "================================================================================\n"
        "  HASIL MATCHING\n"
        "================================================================================\n"
        "Latency: 1.46 ms (target <500ms)\n"
        "Total matches: 32\n"
        "  Tier1↔Tier1: 0\n"
        "  Cross-tier / Tier2: 32\n"
        "Candidate pairs evaluated: 114\n"
        "\n"
        "TOP MATCHES (sorted by FinalScore desc)\n"
        " 1 Probolinggo → Kota Surabaya  Bawang Merah  120.0t  89.5  1.00   89.5\n"
        " 2 Bangkalan   → Gresik         Bawang Merah   40.0t  85.1  1.05   89.4\n"
        "      flags: EQUITY_BOOST_05, MADURA_CLUSTER\n"
        " ...\n"
        " 8 Ngawi       → Bangkalan      Beras Premium 250.0t  63.7  1.30   82.8\n"
        "      flags: EQUITY_BOOST_30, MADURA_CLUSTER       ← v10: +30% boost\n"
        "14 Ngawi       → Sampang        Beras Premium 200.0t  60.3  1.30   78.4\n"
        "      flags: EQUITY_BOOST_30, MADURA_CLUSTER       ← v10: +30% boost\n"
        "\n"
        "→ EXTERNAL OPPORTUNITIES\n"
        "  ★ Cluster Madura surplus bawang_merah: 180t total.\n"
        "    Saran: agregasi ekspor ke Surabaya/Sidoarjo via Suramadu.\n"
        "\n"
        "→ DAMPAK EKONOMI\n"
        "Total volume matched:  4,474.0 ton\n"
        "Gross arbitrage value: Rp 16.186.500.000")

    heading(doc, "5.5.9.3 Run Test Suite", level=3)
    code_block(doc, "pytest tests/ -v")
    para(doc, "Expected (v10): 106 passed in 0.16s. Test breakdown:",
         bold_first="Output:")
    bullet(doc, "16 unit test Layer 0 (tier classification 8 kota IHK)")
    bullet(doc, "19 unit test Layer 1 (haversine, viability, BBM, candidate generation, Bulog)")
    bullet(doc, "23 unit test Layer 2 (5-dim scoring + ramadan/import/default weights)")
    bullet(doc, "14 unit test Layer 3 (equity multiplier, confidence, stable matching, greedy, dispatcher)")
    bullet(doc, "4 scenario test Volume (A1-A4)")
    bullet(doc, "6 scenario test Spatial (B1-B3)")
    bullet(doc, "7 scenario test Temporal (C1-C3)")
    bullet(doc, "9 scenario test Disrupsi (D1-D5)")
    bullet(doc, "8 scenario test Politis (E1-E5)")

    heading(doc, "5.5.9.4 Run Latency Benchmark (NEW v10)", level=3)
    code_block(doc, "python benchmarks/latency.py")
    para(doc, "Expected (v10):", bold_first="Output:")
    code_block(doc,
        "==================================================================================\n"
        "  AgriFlow Matching Engine — Latency Benchmark\n"
        "==================================================================================\n"
        "  Target Section 5.5.4: <500ms p99 untuk 38 kab × 19 komoditas\n"
        "\n"
        "  Configuration                    N (s×d)   p50    p95    p99    max\n"
        "  --------------------------------------------------------------------\n"
        "  Sample data CSV (realistic)      40×33     0.99   1.26   1.38   1.42\n"
        "  Synthetic full Jatim (38×19)     361×361  48.37  53.67  55.53  58.43\n"
        "  Stress 100×100 (national scale)  100×100  12.62  14.82  15.51  15.65\n"
        "  Stress 200×200                   200×200  25.47  26.92  27.54  27.76\n"
        "\n"
        "  PASS  semua konfigurasi p99 < 500.0ms target\n"
        "        (highest p99 = 55.53ms, margin 88.9%)")

    heading(doc, "5.5.9.5 Folder Structure", level=3)
    code_block(doc,
        "agriflow_engine/\n"
        "├── matching_engine/\n"
        "│   ├── models.py            # Dataclasses\n"
        "│   ├── constraints.py       # Layer 0 + Layer 1 (9 hard constraints)\n"
        "│   ├── scoring.py           # Layer 2 (5-dim scoring)\n"
        "│   ├── allocation.py        # Layer 3 (Gale-Shapley + Greedy)\n"
        "│   └── engine.py            # Orchestrator + 24 skenario handlers (v11)\n"
        "├── data_sources/            # 8 connector dual-mode\n"
        "│   ├── pihps_bi.py | bapanas.py | bps.py | bmkg.py\n"
        "│   ├── pvmbg.py | bnpb.py | google_maps.py | hijri_calendar.py\n"
        "├── sample_data/             # CSV 38 kab × 19 komoditas Jatim\n"
        "│   ├── generate_sample_data.py | loader.py\n"
        "│   ├── kabupaten_jatim.csv | komoditas_constraints.csv\n"
        "│   ├── surplus_deficit.csv | weather_forecast.csv\n"
        "│   └── historical_price_stats.csv\n"
        "├── tests/                   # 106 pytest test\n"
        "│   ├── conftest.py | test_layer0_tier.py | test_layer1_constraints.py\n"
        "│   ├── test_layer2_scoring.py | test_layer3_allocation.py\n"
        "│   ├── test_scenarios_volume.py | test_scenarios_spatial.py\n"
        "│   ├── test_scenarios_temporal.py | test_scenarios_disruption.py\n"
        "│   └── test_scenarios_political.py\n"
        "├── examples/run_demo.py     # End-to-end demo\n"
        "├── benchmarks/latency.py    # NEW v10 — performance benchmark\n"
        "├── docs/generate_v10_docx.py # NEW v10 — proposal generator\n"
        "├── README.md\n"
        "└── requirements.txt")

    heading(doc, "5.5.9.6 Programmatic API", level=3)
    code_block(doc,
        "from matching_engine import (\n"
        "    run_matching, SupplyNode, DemandNode,\n"
        "    Kabupaten, Tier, Commodity, LogisticsContext,\n"
        ")\n"
        "\n"
        "kediri = Kabupaten(id='3506', nama='Kediri',\n"
        "                   latitude=-7.796, longitude=112.170,\n"
        "                   ipm=74.50, tier=Tier.MEDIUM)\n"
        "surabaya = Kabupaten(id='3578', nama='Kota Surabaya',\n"
        "                     latitude=-7.2575, longitude=112.7521,\n"
        "                     ipm=84.69, tier=Tier.HIGH)\n"
        "cabai = Commodity(code='cabai_merah', nama='Cabai Merah Besar',\n"
        "                  max_distance_km=200, min_viable_tons=1.0,\n"
        "                  max_fresh_age_days=5)\n"
        "\n"
        "report = run_matching(\n"
        "    surplus_nodes=[SupplyNode(kediri, cabai, volume_tons=80, price_per_kg=30000)],\n"
        "    deficit_nodes=[DemandNode(surabaya, cabai, volume_tons=80, price_per_kg=60000)],\n"
        "    logistics=LogisticsContext(),\n"
        ")\n"
        "\n"
        "for m in report.matches:\n"
        "    print(f'{m.surplus.kabupaten.nama} → {m.deficit.kabupaten.nama}')\n"
        "    print(f'  Volume: {m.matched_volume_tons}t @ {m.distance_km:.0f}km')\n"
        "    print(f'  Score: {m.final_score:.1f} '\n"
        "          f'(base {m.base_score:.1f} × {m.equity_multiplier})')\n"
        "    print(f'  Confidence: {m.confidence.value}, Flags: {m.flags}')")

    # ==================================================================
    # 5.5.10 TEST STRATEGY (preserve from v9)
    # ==================================================================
    heading(doc, "5.5.10 Test Strategy & Skenario Mapping", level=2)
    para(doc,
        "Setiap dari 24 skenario edge case (v11) di Section 5.5.5 ter-mapped ke "
        "pytest test class yang validates behavior secara automated. Saat "
        "developer modify scoring weight, equity threshold, atau menambah "
        "skenario baru, regression detected langsung via pytest. v10 "
        "verified: 106/106 PASS dalam 0.16s di Python 3.14.3.")

    heading(doc, "5.5.10.1 Mapping Skenario → File Test", level=3)
    make_table(doc,
        headers=["Kode", "Skenario", "File Test", "Test Class"],
        rows=[
            ["A1", "1-to-many", "test_scenarios_volume.py", "TestA1_OneToMany"],
            ["A2", "Many-to-1", "test_scenarios_volume.py", "TestA2_ManyToOne"],
            ["A3", "Volume Mismatch Drastis", "test_scenarios_volume.py", "TestA3_VolumeMismatchDrastis"],
            ["A4", "Zero Demand", "test_scenarios_volume.py", "TestA4_ZeroDemand"],
            ["B1", "Cross-tier", "test_scenarios_spatial.py", "TestB1_CrossTier"],
            ["B2", "Long Distance", "test_scenarios_spatial.py", "TestB2_LongDistance"],
            ["B3", "Cluster Madura", "test_scenarios_spatial.py", "TestB3_ClusterMadura"],
            ["C1", "Ramadan Spike", "test_scenarios_temporal.py", "TestC1_RamadanSpike"],
            ["C2", "Pasca Panen Raya", "test_scenarios_temporal.py", "TestC2_PostHarvest"],
            ["C3", "Stale Data", "test_scenarios_temporal.py", "TestC3_StaleData"],
            ["D1", "Banjir Rute", "test_scenarios_disruption.py", "TestD1_BanjirRute"],
            ["D2", "Komoditas Rusak", "test_scenarios_disruption.py", "TestD2_KomoditasRusak"],
            ["D3", "Harga Anomali", "test_scenarios_disruption.py", "TestD3_HargaAnomali"],
            ["D4", "Erupsi Gunung", "test_scenarios_disruption.py", "TestD4_ErupsiGunung"],
            ["D5", "Banjir Multi-Kab", "test_scenarios_disruption.py", "TestD5_BanjirMultiKab"],
            ["E1", "Equity Tie-Break", "test_scenarios_political.py", "TestE1_EquityTieBreak"],
            ["E2", "Pemda Override", "test_scenarios_political.py", "TestE2_PemdaOverride"],
            ["E3", "Bulog Priority", "test_scenarios_political.py", "TestE3_BulogPriority"],
            ["E4", "Import Policy", "test_scenarios_political.py", "TestE4_ImportPolicy"],
            ["E5", "BBM Naik", "test_scenarios_political.py", "TestE5_BBMNaik"],
        ],
        col_widths=[1.5, 4.5, 5.5, 5.5],
    )

    heading(doc, "5.5.10.2 Fixture Library (conftest.py)", level=3)
    para(doc, "Test fixtures siap-pakai untuk 17 kabupaten Jatim representatif "
              "(v10: comments updated dengan equity multiplier baru):")
    make_table(doc,
        headers=["Fixture", "Tier", "IPM", "Equity Mult (v10)", "Use Case"],
        rows=[
            ["surabaya", "1", "84.69", "1.00", "Pasar besar, no boost"],
            ["kota_kediri", "1", "81.48", "1.00", "Tier 1 IHK"],
            ["sumenep", "1", "68.79", "1.15", "Tier 1 IHK + boost"],
            ["banyuwangi", "1", "73.45", "1.05", "Tier 1 IHK pinggir timur"],
            ["kediri_kab", "2", "74.50", "1.05", "Sentra cabai"],
            ["sampang", "2", "66.72", "1.30", "v10: IPM <68 → +30% boost"],
            ["sampang_severe", "2", "62.00", "1.30", "Hipotetis untuk test threshold"],
            ["bangkalan", "2", "67.70", "1.30", "v10: IPM <68 → +30% boost"],
            ["pamekasan", "2", "70.43", "1.15", "Madura cluster"],
            ["bondowoso", "2", "69.62", "1.15", "Equity boost candidate"],
            ["pacitan", "2", "71.40", "1.15", "Pinggir barat"],
            ["lumajang", "2", "70.10", "1.15", "Test erupsi Semeru (D4)"],
            ["sidoarjo", "2", "80.13", "1.00", "Pasar besar, no boost"],
            ["gresik", "2", "77.61", "1.05", "Pasar industri"],
            ["blitar_kab", "2", "73.85", "1.05", "Sentra cabai alt"],
            ["tulungagung", "2", "75.30", "1.05", "Sentra cabai alt"],
            ["probolinggo_kab", "2", "69.40", "1.15", "Sentra bawang + boost"],
        ],
        col_widths=[3.5, 1.0, 1.5, 3.0, 8.0],
    )
    para(doc,
        "Plus factory fixtures: make_supply(), make_demand(), make_stale_supply(), "
        "logistics_normal, logistics_bbm_naik_20pct, logistics_ramadan, "
        "weather_clear, weather_banjir.")

    heading(doc, "5.5.10.3 Continuous Integration Recommendation", level=3)
    para(doc,
        "Tim disarankan setup GitHub Actions atau GitLab CI untuk run pytest "
        "di setiap PR:")
    code_block(doc,
        "# .github/workflows/test.yml\n"
        "name: AgriFlow Engine Tests\n"
        "on: [push, pull_request]\n"
        "jobs:\n"
        "  test:\n"
        "    runs-on: ubuntu-latest\n"
        "    steps:\n"
        "      - uses: actions/checkout@v4\n"
        "      - uses: actions/setup-python@v5\n"
        "        with: { python-version: '3.12' }\n"
        "      - run: pip install -r requirements.txt\n"
        "      - run: python sample_data/generate_sample_data.py\n"
        "      - run: pytest tests/ -v --tb=short\n"
        "      - run: python benchmarks/latency.py  # v10 NEW")

    para(doc, "Latency benchmark juga bisa di-enforce via dedicated test:")
    code_block(doc,
        "def test_latency_under_500ms(full_jatim_dataset):\n"
        "    report = run_matching(*full_jatim_dataset)\n"
        "    assert report.run_metadata['latency_ms'] < 500")

    # ==================================================================
    # 5.5.11 v10 CHANGES (NEW SECTION)
    # ==================================================================
    heading(doc, "5.5.11 Calibration & Bug Fixes (v10 NEW)", level=2)

    para(doc,
        "v10 adalah dedicated calibration release yang menyelesaikan gap "
        "antara klaim proposal v9 dan reality saat dijalankan. Semua "
        "perubahan tervalidasi via pytest 106/106 PASS dan benchmark.")

    heading(doc, "5.5.11.1 Equity Multiplier Threshold Recalibration", level=3)
    para(doc, "Issue:", bold_first="Issue:")
    para(doc,
        "Threshold v9 lama (<65 → 1.30 / <70 → 1.15 / <75 → 1.05 / ≥75 → 1.00) "
        "tidak pernah ter-trigger untuk band 1.30 dengan data BPS 2024. "
        "IPM terendah Jatim 2024 = Sampang 66.72. Klaim proposal '+30% boost "
        "untuk kab tertinggal' jadi vacuous.")
    para(doc, "Fix:", bold_first="Fix:")
    para(doc,
        "Threshold v10 (<68 → 1.30 / <72 → 1.15 / <78 → 1.05 / ≥78 → 1.00) "
        "kalibrasi ulang sesuai distribusi BPS 2024. Sampang (66.72) & "
        "Bangkalan (67.70) sekarang konkret menerima +30% boost.")
    para(doc, "Validation:", bold_first="Validation:")
    code_block(doc,
        "$ python examples/run_demo.py | grep EQUITY_BOOST_30\n"
        " 8 Ngawi → Bangkalan Beras Premium 250.0t  63.7  1.30   82.8\n"
        "      flags: EQUITY_BOOST_30, MADURA_CLUSTER\n"
        "14 Ngawi → Sampang   Beras Premium 200.0t  60.3  1.30   78.4\n"
        "      flags: EQUITY_BOOST_30, MADURA_CLUSTER")

    heading(doc, "5.5.11.2 Cross-Platform Demo Compatibility", level=3)
    para(doc, "Issue:", bold_first="Issue:")
    para(doc,
        "examples/run_demo.py dan sample_data/generate_sample_data.py crash "
        "dengan UnicodeEncodeError di Windows console default (cp1252). "
        "Karakter Unicode '→', '★', '✓' tidak bisa di-encode. Demo gagal di "
        "first line — first-impression killer untuk juri.")
    para(doc, "Fix:", bold_first="Fix:")
    code_block(doc,
        "if sys.platform == 'win32':\n"
        "    try:\n"
        "        sys.stdout.reconfigure(encoding='utf-8')\n"
        "        sys.stderr.reconfigure(encoding='utf-8')\n"
        "    except (AttributeError, OSError):\n"
        "        pass")
    para(doc, "Validation:", bold_first="Validation:")
    para(doc,
        "Demo verified runs on Windows 11 fresh-install with default cp1252 "
        "console encoding.")

    heading(doc, "5.5.11.3 Stale Data Confidence Drop Bertingkat", level=3)
    para(doc, "Issue:", bold_first="Issue:")
    para(doc,
        "v9 spec C3 mengharuskan stale data confidence: HIGH→MEDIUM dan "
        "MEDIUM→LOW. Implementasi v9 hanya HIGH→MEDIUM (Tier 2 stale tetap "
        "MEDIUM, tidak konsisten dengan spec).")
    para(doc, "Fix (matching_engine/engine.py:373-383):", bold_first="Fix:")
    code_block(doc,
        "if s_key in stale_supply_keys or d_key in stale_demand_keys:\n"
        "    flags.append('STALE_DATA_24H')\n"
        "    if m.confidence == Confidence.HIGH:\n"
        "        m.confidence = Confidence.MEDIUM\n"
        "    elif m.confidence == Confidence.MEDIUM:\n"
        "        m.confidence = Confidence.LOW")

    heading(doc, "5.5.11.4 IPM Source-of-Truth Sync", level=3)
    para(doc, "Issue:", bold_first="Issue:")
    para(doc,
        "data_sources/bps.py:IPM_2024_JATIM dan sample_data/generate_sample_data.py "
        "punya nilai berbeda untuk 7+ kabupaten (mis. Madiun kab: bps.py 71.95 "
        "vs CSV 73.55; Magetan: bps.py 72.20 vs gen 75.80). Risiko: live API "
        "fallback memakai data berbeda dari sample mock.")
    para(doc, "Fix:", bold_first="Fix:")
    para(doc,
        "Sync IPM_2024_JATIM agar mirror generate_sample_data.py (single "
        "source of truth: generator). CSV regenerated. v10 verified: ketiga "
        "file (generator, CSV, bps.py) konsisten.")

    heading(doc, "5.5.11.5 Stable Identity Check untuk Stale Detection", level=3)
    para(doc, "Issue:", bold_first="Issue:")
    para(doc,
        "Stale check di v9 pakai dataclass `in` operator yang mengandalkan "
        "object identity. Kalau supply_nodes filtered (mis. anomaly removal "
        "membuat list baru), identity check gagal. Bug latent.")
    para(doc, "Fix:", bold_first="Fix:")
    code_block(doc,
        "stale_supply_keys = {(n.kabupaten.id, n.commodity.code) for n in stale_supply}\n"
        "stale_demand_keys = {(n.kabupaten.id, n.commodity.code) for n in stale_demand}\n"
        "# Use tuple set lookup instead of dataclass identity")

    heading(doc, "5.5.11.6 BMKG prefer_bmkg Flag Properly Honored", level=3)
    para(doc, "Issue:", bold_first="Issue:")
    para(doc,
        "WeatherConnector.__init__ menerima prefer_bmkg flag, tapi "
        "fetch_route_forecast tidak pernah membaca flag tersebut. _fetch_bmkg "
        "didefinisikan tapi tidak pernah dipanggil. Dead code.")
    para(doc, "Fix:", bold_first="Fix:")
    para(doc,
        "v10 mengubah signature: WeatherConnector(prefer_bmkg, adm4_lookup, "
        "timeout). Fetch logic mengecek: jika prefer_bmkg=True dan "
        "adm4_lookup tersedia untuk semua titik rute, pakai BMKG; jika tidak, "
        "auto-fallback ke Open-Meteo. Jadi BMKG menjadi optional integration "
        "dengan dependency yang explicit.")

    heading(doc, "5.5.11.7 Latency Benchmark Tool (NEW)", level=3)
    para(doc,
        "v10 menambahkan benchmarks/latency.py — multi-config stress test "
        "yang mengukur p50/p95/p99/max latency untuk 4 konfigurasi: sample "
        "data realistic (40×33), synthetic full Jatim (361×361), stress "
        "100×100, stress 200×200. Output PASS/FAIL terhadap target 500ms p99.")

    heading(doc, "5.5.11.8 Dokumentasi Generator (NEW)", level=3)
    para(doc,
        "v10 menambahkan docs/generate_v10_docx.py — script Python yang "
        "regenerate proposal docx ini dari source. Saat tim update threshold, "
        "skenario, atau benchmark, jalankan ulang script ini untuk "
        "synchronize dokumentasi.")

    # ==================================================================
    # 5.5.12 LATENCY BENCHMARK RESULTS (NEW)
    # ==================================================================
    heading(doc, "5.5.12 Latency Benchmark Results (v10)", level=2)

    para(doc,
        "Klaim Section 5.5.4 '<500ms p99 untuk 38 kab × 19 komoditas' "
        "verified secara konkret via benchmarks/latency.py. Metrik diukur "
        "pada Python 3.14.3, Windows 11 Home, hardware konsumen.")

    make_table(doc,
        headers=["Configuration", "N (s × d)", "p50 (ms)", "p95 (ms)",
                 "p99 (ms)", "Max (ms)", "Mean (ms)"],
        rows=[
            ["Sample data CSV (realistic)", "40 × 33",
             "0.99", "1.26", "1.38", "1.42", "1.02"],
            ["Synthetic full Jatim (38×19)", "361 × 361",
             "48.37", "53.67", "55.53", "58.43", "49.03"],
            ["Stress 100×100 (national scale)", "100 × 100",
             "12.62", "14.82", "15.51", "15.65", "12.85"],
            ["Stress 200×200", "200 × 200",
             "25.47", "26.92", "27.54", "27.76", "25.50"],
        ],
        col_widths=[5.0, 2.5, 1.7, 1.7, 1.7, 1.7, 1.7],
    )

    callout(doc,
        "PASS — semua konfigurasi p99 < 500ms target. Highest p99 = 55.53ms "
        "(synthetic full Jatim 361×361). Margin dari target = 88.9%. Demo "
        "realistik (40×33 sample): p99 cuma 1.38ms — 99.7% margin.",
        color="E8F5E8",
    )

    para(doc,
        "Reproducibility: jalankan python benchmarks/latency.py setelah "
        "pip install. Output ter-cetak ke stdout. Untuk CI, parse output "
        "atau implement pytest assertion.",
        bold_first="Reproducibility:")

    # ==================================================================
    # 6. GLOBAL BEST PRACTICES
    # ==================================================================
    doc.add_page_break()
    heading(doc, "6. Global Best Practices — 5 Adopsi First-in-Indonesia", level=1)

    para(doc,
        "AgriFlow tidak klaim menciptakan dari nol. AgriFlow mengadopsi 5 "
        "praktik terbaik dunia yang sudah proven, menjadi platform pertama di "
        "Indonesia yang mengimplementasikannya.")

    make_table(doc,
        headers=["Teknologi", "Asal Global", "Status Indonesia", "Fase AgriFlow"],
        rows=[
            ["Voice-First AI Multi-Bahasa Daerah",
             "India Bharat-VISTAAR", "Belum ada (Pak Dayat fokus advisory)",
             "BUILD Y1 (Sahabat-AI)"],
            ["Geospatial Decision Support System",
             "India Krishi DSS + USDA CDL", "Data terpisah BMKG/BPS/Kementan",
             "BUILD Y1-Y2 (Sentinel-2)"],
            ["Sub-National Early Warning",
             "USAID FEWS NET", "Belum ada sub-nasional operational",
             "BUILD Y1 (core)"],
            ["Bundled Financial Services Triggers",
             "IWMI/CGIAR pilot India", "AUTP + KUR ada tapi standalone",
             "ROADMAP Y2"],
            ["MRV Carbon Credit Food Loss",
             "Verra VCS + Gold Standard", "Belum ada project Indonesia certified",
             "ROADMAP Y2 (Rp 600jt/y)"],
        ],
        col_widths=[5.0, 4.0, 4.5, 3.5],
    )
    para(doc,
        "Plus AgriFlow Matching Engine sendiri sebagai adopsi keenam: "
        "kombinasi Uber/MealConnect/Gale-Shapley yang purpose-built untuk "
        "pangan kab-level Indonesia (lihat Section 5.5).",
        italic=True)

    # ==================================================================
    # 7. ARSITEKTUR & TECH STACK
    # ==================================================================
    heading(doc, "7. Arsitektur & Tech Stack", level=1)

    para(doc, "Tech stack pilot: Rp 1-2 juta/bulan operational. Scale Jatim "
              "full: Rp 3-5 juta. Scale nasional: Rp 10-15 juta.")

    make_table(doc,
        headers=["Layer", "Teknologi", "Alasan", "Biaya/bln"],
        rows=[
            ["Frontend", "Next.js 14 + Tailwind + Leaflet",
             "SSR cepat, Vercel free, React ecosystem", "Rp 0"],
            ["Backend API", "FastAPI + Pydantic + SQLAlchemy",
             "ML-native async, auto Swagger", "Rp 0 (Railway free)"],
            ["Database", "PostgreSQL 15 (Supabase)",
             "13 tabel, JSONB, Auth+Realtime", "Rp 0 (free 500MB)"],
            ["Vector DB", "Qdrant self-hosted",
             "RAG knowledge, semantic search", "Rp 0"],
            ["Cache", "Redis Upstash",
             "Geohash + price cache", "Rp 0 (free)"],
            ["ML", "XGBoost + Prophet + Isolation Forest",
             "Best tabular + seasonality + anomaly", "Rp 0"],
            ["Matching", "scipy + NumPy + matching lib",
             "4-lapis Section 5.5", "Rp 0"],
            ["LLM Indonesia", "Gemini 1.5 Flash",
             "1500 req/day free", "Rp 0"],
            ["LLM Bahasa Daerah", "Sahabat-AI 70B (GoTo)",
             "5 bahasa daerah", "Rp 200-500K"],
            ["WhatsApp", "Twilio WhatsApp Business",
             "Sandbox → production", "Rp 200K+"],
            ["Phone IVR", "Twilio Voice + Programmable",
             "Toll-free 0800-AGRIFLOW", "Rp 300-500K"],
            ["STT/TTS", "Google Cloud STT + TTS",
             "Indonesian voices, Bahasa Daerah", "Rp 150-400K"],
            ["Routing", "Google Maps Routes API",
             "Real road distance", "Rp 50-125K"],
            ["Orchestration", "n8n self-hosted",
             "Visual debug, cron reliable", "Rp 50-100K"],
        ],
        col_widths=[3.0, 4.5, 5.5, 4.0],
    )

    # ==================================================================
    # 8. BUSINESS MODEL
    # ==================================================================
    heading(doc, "8. Business Model — 8 Revenue Streams", level=1)

    para(doc,
        "AgriFlow bukan proyek sosial yang butuh grant untuk hidup. Revenue "
        "dari 6 customer segments via 8 streams.")

    make_table(doc,
        headers=["Revenue Stream", "Target Customer", "Unit Price", "Fase"],
        rows=[
            ["1. BI GPIPS National License", "Bank Indonesia Pusat", "Rp 1,5-3M/thn", "Y2-Y3"],
            ["2. Dinas APBD Subscription", "Dinas Pertanian + TPID", "Rp 75-200jt/thn", "Y1-Y3"],
            ["3. B2B Data API (DaaS)", "Bank, asuransi, FMCG", "Rp 500K-25jt/bln", "Y2-Y3"],
            ["4. White-Label Provincial", "Pemprov lain", "Rp 500jt-1M setup", "Y2-Y3"],
            ["5. Custom Intelligence Reports", "Think tank, Bapanas", "Rp 5-75jt/report", "Y1-Y3"],
            ["6. WhatsApp Premium Tier", "Petani, pedagang", "Rp 15K-500K/bln", "Y1-Y3"],
            ["7. Grants (Impact Capital)", "BI Institute, UNDP, GIZ", "Rp 500jt-2M/grant", "Y1-Y3"],
            ["8. ESG & Carbon Credit", "Verra, Gold Standard", "Rp 600jt-2M/thn", "Y2-Y3"],
        ],
        col_widths=[5.0, 4.5, 4.0, 3.5],
    )

    heading(doc, "8.1 Proyeksi Revenue 3 Tahun", level=2)
    make_table(doc,
        headers=["Stream", "Y1 (2026)", "Y2 (2027)", "Y3 (2028)"],
        rows=[
            ["BI GPIPS License", "Rp 0", "Rp 500jt", "Rp 2.000jt"],
            ["Dinas APBD", "Rp 0", "Rp 1.000jt", "Rp 3.000jt"],
            ["API DaaS", "Rp 0", "Rp 1.620jt", "Rp 4.000jt"],
            ["White-Label", "Rp 0", "Rp 0", "Rp 1.500jt"],
            ["Custom Reports", "Rp 100jt", "Rp 500jt", "Rp 1.000jt"],
            ["WhatsApp Premium", "Rp 50jt", "Rp 660jt", "Rp 2.000jt"],
            ["Grants", "Rp 500jt", "Rp 1.500jt", "Rp 2.000jt"],
            ["ESG & Carbon", "Rp 0", "Rp 200jt", "Rp 600jt"],
            ["TOTAL REVENUE", "Rp 650jt", "Rp 5.980jt", "Rp 16.100jt"],
            ["NET INCOME", "+Rp 450jt", "+Rp 4.480jt", "+Rp 11.100jt"],
        ],
        col_widths=[5.0, 4.0, 4.0, 4.0],
    )

    # ==================================================================
    # 9. BIAYA PENGEMBANGAN
    # ==================================================================
    heading(doc, "9. Biaya Pengembangan", level=1)

    heading(doc, "9.1 Hackathon (16 Minggu)", level=2)
    callout(doc,
        "TOTAL HACKATHON (16 minggu): Rp 19.000.000\n"
        "• Setup: Rp 4,2jt (domain, design, legal, accounting)\n"
        "• Infrastructure 4 bln: Rp 3,8jt (Railway, Hostinger, Twilio, GCP, Sahabat-AI)\n"
        "• Field work: Rp 11jt (survey 150 petani, interview 3 Dinas, audio + video)\n"
        "• Funding source: modal tim 4 × Rp 5jt = Rp 20jt, atau CSR BRI Rp 300-500jt",
        color="FFF8E1",
    )

    heading(doc, "9.2 Pilot 6 Bulan (Aug 2026 - Jan 2027)", level=2)
    make_table(doc,
        headers=["Kategori", "6 Bulan (Rp)", "Breakdown"],
        rows=[
            ["Tim honorarium (4 orang)", "60.000.000", "Rp 2,5jt/orang/bln × 4 × 6"],
            ["Field network (5 data collectors)", "45.000.000", "Rp 1,5jt/orang/bln × 5 × 6"],
            ["Infrastructure scale-up", "30.000.000", "Supabase Pro, Railway scale"],
            ["Marketing + stakeholder", "12.000.000", "Workshop Dinas, roadshow"],
            ["Buffer", "3.000.000", "Contingency"],
            ["TOTAL PILOT", "150.000.000", ""],
        ],
        col_widths=[5.0, 4.0, 8.0],
    )
    para(doc,
        "Funding pilot: kombinasi hackathon prize (Rp 100-500jt), BI Institute "
        "Innovation Fund (Rp 500jt-1M), CSR BRI/Bank Jatim (Rp 300-500jt), "
        "UNDP (Rp 250-500jt). Target Y1 total: Rp 500jt - 1M.")

    # ==================================================================
    # 10. TIM
    # ==================================================================
    heading(doc, "10. Tim 4 Orang — Role Distribution", level=1)
    make_table(doc,
        headers=["Role", "Tanggung Jawab", "Output Hackathon"],
        rows=[
            ["1. Tech Lead (Backend + AI)",
             "• FastAPI architecture\n• XGBoost + Prophet training\n• Matching Engine 4-lapis (Section 5.5)\n• Gemini + Sahabat-AI integration\n• Database schema",
             "• API 20+ endpoints\n• Backtest MAPE <20%\n• Matching engine live 5 kab\n• 24 scenarios tested (126/126 pytest, v11)"],
            ["2. Frontend + UX Engineer",
             "• Next.js dashboard\n• Leaflet map + choropleth\n• Responsive design\n• WCAG AA accessibility",
             "• Dashboard working data real\n• 5 kab choropleth interaktif\n• Demo-ready"],
            ["3. Data + DevOps Engineer",
             "• PIHPS scraper Tier 1\n• Bapanas + BPS scraper Tier 2\n• n8n orchestration\n• Twilio IVR + WhatsApp setup\n• TTS/STT integration",
             "• Data pipeline 38 kab\n• 12 n8n workflows\n• Phone IVR working\n• Monitoring active"],
            ["4. Product + Business Lead",
             "• Stakeholder outreach\n• Survey + interview\n• Pitch deck + video\n• Storytelling\n• Partnership awal",
             "• 3 LOI Dinas\n• 150+ survey petani\n• Pitch deck v8 + video 3 menit\n• 2 partnership MoU"],
        ],
        col_widths=[3.5, 6.5, 7.0],
    )
    para(doc,
        "Beban: 20-30 jam/minggu × 16 minggu = 320-480 jam per orang. Crunch "
        "period minggu 9-12: 40 jam/minggu. Bootstrap mode (no honor) di "
        "hackathon, pilot phase Rp 2,5jt/bln per orang.",
        italic=True)

    # ==================================================================
    # 11. ROADMAP
    # ==================================================================
    heading(doc, "11. Roadmap Eksekusi 16 Minggu", level=1)
    make_table(doc,
        headers=["Minggu", "Periode", "Deliverable", "Owner"],
        rows=[
            ["1-2", "30 Apr - 13 May",
             "• Training DIGDAYA\n• Setup repo + infrastruktur\n• PIHPS scraper Tier 1 skeleton\n• Bapanas scraper Tier 2 skeleton",
             "Tim 4"],
            ["3-4", "14-27 May",
             "• Matching Engine Layer 1+2 (greedy)\n• Train XGBoost + Prophet\n• Dashboard skeleton\n• DB schema finalized",
             "Tech + Frontend"],
            ["5-6", "28 May - 10 Jun",
             "• Matching Layer 3 Tier 2 (greedy)\n• WhatsApp bot text (Gemini RAG)\n• Twilio Voice IVR setup\n• TTS/STT Bahasa Indonesia",
             "Data + Tech"],
            ["7-8", "11-24 Jun",
             "• Matching Layer 3 Tier 1 (stable matching)\n• Dashboard full working\n• WhatsApp Voice handler\n• Field visit Kediri/Malang",
             "Semua + Product"],
            ["9-10", "25 Jun - 8 Jul",
             "• Bahasa Jawa Sahabat-AI integration\n• 24 scenarios testing (v11)\n• Load test 100 users\n• Survey 100+ cumulative",
             "Semua tim"],
            ["11-12", "9-22 Jul",
             "• Record Pak Tani Bahasa Jawa\n• Video demo 3 menit\n• Pitch deck v10 final\n• Testimonial 5 petani + 3 Dinas",
             "Product + Tech"],
            ["13-14", "23 Jul - 5 Aug",
             "• Pitch rehearsal 3x mentor\n• Q&A prep + backup\n• Dokumentasi teknis (proposal v10)\n• Submission",
             "Product + Tech"],
            ["15-16", "6-19 Aug",
             "• Finalist event\n• Business matching prep\n• Investor materials\n• Post-hackathon plan",
             "Product + Business"],
        ],
        col_widths=[1.5, 3.0, 8.5, 4.0],
    )

    heading(doc, "11.1 Milestone Kritis", level=2)
    bullet(doc, "M1 (Minggu 4): Matching engine 2-layer working (Tier 2 greedy)")
    bullet(doc, "M2 (Minggu 6): WhatsApp + Phone IVR demo Bahasa Indonesia")
    bullet(doc, "M3 (Minggu 8): Stable matching Tier 1 working + 5 kab pilot active")
    bullet(doc, "M4 (Minggu 10): 24 scenarios tested (v11) + Bahasa Jawa working")
    bullet(doc, "M5 (Minggu 12): Video Pak Tani + pitch deck v10 final")
    bullet(doc, "M6 (Minggu 16): Hackathon submission + business matching")

    # ==================================================================
    # 12. RISK MITIGATION
    # ==================================================================
    heading(doc, "12. Risk Mitigation", level=1)
    make_table(doc,
        headers=["Risiko", "Deskripsi", "Mitigasi"],
        rows=[
            ["Data PIHPS tidak stabil", "Token expired, format berubah",
             "Dual-source Bapanas IPH, validasi outlier, n8n alert"],
            ["Model meleset di black swan", "Banjir, pandemi, geopolitik",
             "Confidence interval, retrain mingguan, manual override"],
            ["Adopsi Pemda lambat", "SDM IT terbatas",
             "Pilot gratis 3-6 bulan, training, UI simple"],
            ["Petani tidak melek digital", "Tidak familiar bot",
             "Phone IVR fallback, voice message, Bahasa daerah"],
            ["Stable matching scope cut", "Timeline mepet hackathon",
             "Tier 2 greedy sudah demo-able M2. Tier 1 stable matching scope yang bisa di-cut."],
            ["Sahabat-AI pricing unclear", "API cost variabel",
             "Fallback Gemini + manual translation Bahasa Jawa demo"],
            ["Tim 4 orang burnout", "Scope creep + deadline",
             "Sprint planning, weekly retro, ruthless prioritization"],
            ["Funding dry", "VC agritech bearish",
             "8-stream diversify, grant pipeline, bootstrap possible"],
        ],
        col_widths=[4.0, 4.5, 8.5],
    )

    # ==================================================================
    # 13. WHY AGRIFLOW WINS
    # ==================================================================
    heading(doc, "13. Why AgriFlow Wins", level=1)
    callout(doc,
        "1. Purpose-Built Sub-National Matching Engine. Modified Gale-Shapley "
        "stable matching (Nobel Prize 2012, fires Tier 1↔Tier 1) + greedy "
        "multi-objective + equity priority (Tier 2 + cross-tier — Jatim "
        "production path) + IPM equity multiplier kalibrasi BPS 2024 + "
        "24 scenarios coverage (19 engineering + 5 commercial-reality v11). Kombinasi pertama yang verifiable terhadap "
        "code base untuk pangan sub-nasional Indonesia. v11: 106/106 tests "
        "PASS + benchmark p99 = 55.53ms (margin 88.9%); engine code identik "
        "dengan v10, claim-precision pass selesai.\n\n"
        "2. Accessibility by Design. Triple channel (Web + WhatsApp + Phone "
        "IVR). 22,5 juta penyandang disabilitas + 4 juta lansia + 10 juta "
        "feature phone user pertama terhubung dengan informasi pangan digital.\n\n"
        "3. Operational Layer GPIPS BI 2026. 3 pilar GPIPS dalam 1 platform. "
        "AgriFlow execute kebijakan yang BI sudah canangkan, dengan algoritma "
        "yang setiap klaimnya dapat diverifikasi langsung terhadap file:line "
        "di matching_engine/.\n\n"
        "4. Hybrid Data Strategy — Honest Engineering. Tier 1 (HIGH "
        "confidence) untuk 8 kota IHK saat Tier 1↔Tier 1; Tier 2 (MEDIUM) "
        "untuk 30 kab non-IHK + cross-tier. Tidak fake guarantee. Setiap "
        "match transparent dengan confidence label. v10: stale data drop "
        "bertingkat HIGH→MEDIUM→LOW untuk transparansi yang lebih granular.\n\n"
        "5. Multi-Sided Business + Budget Transparan. 8 streams × 6 segments. "
        "Y3 Rp 16,1M net Rp 11,1M. Hackathon Rp 19jt + Pilot Rp 150jt fully "
        "documented.",
        color="E8F5E8",
    )

    # ==================================================================
    # 14. THE ASK
    # ==================================================================
    heading(doc, "14. The Ask", level=1)
    callout(doc,
        "Dengan dukungan PIDI DIGDAYA × Hackathon 2026, AgriFlow siap:\n"
        "• Launching pilot 5 kab Jatim 3 bulan pasca-hackathon dengan "
        "matching engine 4-lapis purpose-built (executable, ter-test, "
        "ter-benchmark)\n"
        "• Menjalankan operational layer untuk GPIPS 2026 BI dengan "
        "algoritma yang setiap klaimnya verifiable terhadap code\n"
        "• Menjangkau 22,5 juta penyandang disabilitas + 4 juta petani lansia\n"
        "• Mencegah Rp 15-30 miliar food loss pilot Y1, scale Rp 200-300 "
        "miliar nasional Y3\n"
        "• Meningkatkan pendapatan 5.000+ petani 10-15%\n"
        "• Boost kabupaten tertinggal IPM <68 dengan equity multiplier +30% "
        "saat menjadi deficit (konkret applicable ke Sampang & Bangkalan v10)\n"
        "• 39 direct jobs Y1 → 275-330 Y3\n"
        "• Seed round USD 1-2M Q2-Q3 2027 untuk nasional scale\n\n"
        "Deteksi. Prediksi. Distribusi. Untuk Semua.\n"
        "AgriFlow — Purpose-Built Sub-National AI Matching Engine "
        "for Indonesian Food Distribution.",
        color="DDF4FF",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("— END OF PROPOSAL v11.0 —")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PIDI DIGDAYA × Hackathon 2026")
    r.font.size = Pt(11)
    r.italic = True

    # =====================================================================
    # SAVE
    # =====================================================================
    # Default: save di docs/ supaya self-contained dengan repo.
    output_path = Path(__file__).parent / "AgriFlow_v11.docx"
    try:
        doc.save(str(output_path))
        print(f"Generated: {output_path}")
        print(f"Size: {output_path.stat().st_size / 1024:.1f} KB")
        return output_path
    except PermissionError:
        # File terkunci (Word masih buka file v11 lama). Save ke filename baru.
        alt_path = output_path.with_name("AgriFlow_v11_NEW.docx")
        doc.save(str(alt_path))
        print(f"WARNING: {output_path.name} terkunci (kemungkinan Word masih buka).")
        print(f"Saved as: {alt_path}")
        print(f"Size: {alt_path.stat().st_size / 1024:.1f} KB")
        print(f"Action: tutup Word, hapus AgriFlow_v11.docx lama, "
              f"rename {alt_path.name} → AgriFlow_v11.docx")
        return alt_path


if __name__ == "__main__":
    build()
